"""Generate SalonBazar Technical Architecture Document as DOCX."""

from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "docs/Technical-Architecture-SalonBazar.docx"


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1A1714")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # COVER
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SalonBazar")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(26, 23, 20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Technical Architecture Document")
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(201, 168, 76)

    for line in [
        "Version: 1.0",
        f"Date: {date.today().strftime('%B %d, %Y')}",
        "System: SalonBazar Platform",
        "Production URL: https://salonbazar.shop",
        "API Base URL: https://api.salonbazar.shop/api/v1",
        "Audience: Engineering, DevOps, Security, Technical Leadership",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # TOC
    add_heading(doc, "Table of Contents", 1)
    toc = [
        "1. Document Control",
        "2. Purpose & Scope",
        "3. Architecture Principles & Goals",
        "4. System Context (C4 Level 1)",
        "5. High-Level Architecture",
        "6. Application Layer Architecture",
        "7. Client Application Architecture",
        "8. Backend API Architecture",
        "9. Data Architecture",
        "10. Authentication & Authorization",
        "11. Booking Flow Architecture",
        "12. Payment Architecture (Razorpay)",
        "13. Queue & Real-Time Architecture",
        "14. State Management Architecture",
        "15. Service Integration Layer",
        "16. Media & File Upload Architecture",
        "17. API Design Standards",
        "18. Security Architecture",
        "19. Infrastructure & Deployment",
        "20. Environment Configuration",
        "21. Observability & Monitoring",
        "22. Performance & Scalability",
        "23. Error Handling & Resilience",
        "24. CI/CD Pipeline",
        "25. Current vs Target Architecture",
        "26. Technology Stack Reference",
        "27. Architecture Decision Records (ADRs)",
        "28. Appendices",
    ]
    add_numbered(doc, toc)
    doc.add_page_break()

    # 1. DOCUMENT CONTROL
    add_heading(doc, "1. Document Control", 1)
    add_table(doc, ["Field", "Value"], [
        ["Title", "SalonBazar Technical Architecture Document"],
        ["Version", "1.0"],
        ["Status", "Draft for Review"],
        ["Related PRD", "docs/PRD-SalonBazar.docx"],
        ["Owner", "Engineering Team"],
        ["Last Updated", date.today().strftime("%Y-%m-%d")],
    ], [2.0, 4.5])

    add_heading(doc, "1.1 Revision History", 2)
    add_table(doc, ["Version", "Date", "Author", "Changes"], [
        ["1.0", date.today().strftime("%Y-%m-%d"), "Engineering", "Initial technical architecture document"],
    ], [0.7, 1.0, 1.3, 3.5])

    # 2. PURPOSE
    add_heading(doc, "2. Purpose & Scope", 1)
    add_para(doc,
        "This document describes the technical architecture of the SalonBazar platform — a salon discovery, "
        "appointment booking, walk-in queue management, and salon operations system. It is intended for "
        "engineers, architects, DevOps, and security reviewers responsible for building, deploying, and "
        "operating the platform.")
    add_heading(doc, "2.1 In Scope", 2)
    add_bullets(doc, [
        "Web client application architecture (customer, salon owner, admin interfaces)",
        "REST API backend architecture and endpoint design",
        "Authentication, authorization, and session management",
        "Booking, payment, and queue subsystems",
        "Data model, storage, and real-time messaging",
        "Third-party integrations (Razorpay, Supabase, Cloudinary, Google OAuth)",
        "Deployment, environments, security, and operational concerns",
    ])
    add_heading(doc, "2.2 Out of Scope", 2)
    add_bullets(doc, [
        "Detailed UI/UX specifications (see PRD)",
        "Native mobile app architecture (Phase 2)",
        "Business requirements and user stories (see PRD)",
        "Project management timelines and sprint planning",
    ])

    # 3. PRINCIPLES
    add_heading(doc, "3. Architecture Principles & Goals", 1)
    add_table(doc, ["Principle", "Description", "Rationale"], [
        ["Separation of Concerns", "Client, API, and data layers are independently deployable", "Enables parallel development and scaling"],
        ["API-First Design", "All business logic exposed via versioned REST API", "Supports web now, mobile apps later"],
        ["Stateless Backend", "API servers hold no session state; JWT for auth", "Horizontal scaling without sticky sessions"],
        ["Secure by Default", "HTTPS, JWT, RBAC, server-side payment verification", "Protects PII and payment data"],
        ["Real-Time Where Needed", "Supabase Realtime for queue; REST for CRUD", "Low-latency queue without over-engineering"],
        ["Progressive Enhancement", "Mock data during UI development; API swap at integration", "Faster frontend iteration"],
        ["Mobile-First Client", "Responsive SPA optimized for mobile browsers", "Primary user device in target market"],
        ["Fail Gracefully", "Toast notifications, retry logic, token refresh", "Better UX during network failures"],
    ], [1.5, 2.5, 2.3])

    add_heading(doc, "3.1 Architecture Quality Goals", 2)
    add_table(doc, ["Quality Attribute", "Target"], [
        ["Availability", "99.5% uptime for API and web client"],
        ["Latency (API p95)", "< 500ms for read operations"],
        ["Latency (API p95)", "< 800ms for write operations (booking, payment)"],
        ["Concurrent users", "10,000+ simultaneous web sessions"],
        ["Data durability", "Zero data loss for confirmed bookings and payments"],
        ["Recovery Time (RTO)", "< 1 hour for production incidents"],
        ["Recovery Point (RPO)", "< 5 minutes for transactional data"],
    ], [2.0, 4.5])

    # 4. SYSTEM CONTEXT
    add_heading(doc, "4. System Context (C4 Level 1)", 1)
    add_para(doc, "SalonBazar sits at the center of an ecosystem connecting customers, salon owners, and platform administrators with external services.")
    add_table(doc, ["Actor / System", "Type", "Interaction"], [
        ["Customer", "Human", "Discovers salons, books appointments, joins queue, pays online"],
        ["Salon Owner", "Human", "Manages bookings, queue, services, staff via dashboard"],
        ["Platform Admin", "Human", "Onboards salons, moderates users, views platform analytics"],
        ["SalonBazar Web App", "Software", "Single-page application served at salonbazar.shop"],
        ["SalonBazar API", "Software", "REST API at api.salonbazar.shop/api/v1"],
        ["Razorpay", "External", "Payment gateway for UPI, cards, net banking"],
        ["Supabase", "External", "Real-time Postgres changes for live queue updates"],
        ["Cloudinary", "External", "Image CDN for salon gallery, avatars, staff photos"],
        ["Google OAuth", "External", "Social authentication provider"],
        ["SMS/Email Provider", "External", "OTP, booking confirmations, reminders (planned)"],
    ], [1.5, 1.0, 3.8])

    add_heading(doc, "4.1 Context Diagram (Textual)", 2)
    add_code_block(doc, """
                    ┌─────────────────┐
                    │    Customer     │
                    └────────┬────────┘
                             │ HTTPS
    ┌──────────────┐    ┌──────▼──────┐    ┌─────────────────┐
    │ Salon Owner  │───►│  SalonBazar │◄───│ Platform Admin  │
    └──────────────┘    │   Web SPA   │    └─────────────────┘
                        └──────┬──────┘
                               │ REST + JWT
                        ┌──────▼──────┐
                        │ SalonBazar  │
                        │  REST API   │
                        └──┬───┬───┬──┘
              ┌────────────┘   │   └────────────┐
       ┌──────▼─────┐   ┌─────▼─────┐   ┌──────▼──────┐
       │ PostgreSQL │   │ Supabase  │   │  Razorpay   │
       │  (Primary) │   │ Realtime  │   │  Payments   │
       └────────────┘   └───────────┘   └─────────────┘
              ┌──────────────┐   ┌──────────────┐
              │  Cloudinary  │   │ Google OAuth │
              │     CDN      │   │              │
              └──────────────┘   └──────────────┘
    """)

    # 5. HIGH-LEVEL
    add_heading(doc, "5. High-Level Architecture", 1)
    add_para(doc, "SalonBazar follows a three-tier, API-centric architecture:")
    add_numbered(doc, [
        "Presentation Tier — React SPA (Vite build) served via CDN/static hosting",
        "Application Tier — Node.js/Python REST API (api.salonbazar.shop) with business logic",
        "Data Tier — PostgreSQL primary database; Supabase for realtime; Cloudinary for media",
    ])

    add_heading(doc, "5.1 Tier Responsibilities", 2)
    add_table(doc, ["Tier", "Responsibility", "Technology"], [
        ["Presentation", "UI rendering, routing, client state, form validation, Razorpay checkout modal", "React 18, Vite 5, Redux Toolkit, React Router 6"],
        ["Application", "Auth, booking logic, slot management, payments, queue ops, admin", "REST API (Node.js/Express or similar)"],
        ["Data", "Persistent storage, transactions, indexes, backups", "PostgreSQL 15+"],
        ["Real-Time", "Live queue position broadcasts, booking status updates", "Supabase Realtime (Postgres CDC)"],
        ["Media", "Image upload, resize, CDN delivery", "Cloudinary (via API proxy)"],
        ["Payments", "Order creation, signature verification, refunds", "Razorpay (server-side verify)"],
    ], [1.2, 3.0, 2.1])

    add_heading(doc, "5.2 Request Flow (Typical Booking)", 2)
    add_numbered(doc, [
        "Customer selects service, date, slot in Web SPA → state stored in Redux booking slice",
        "SPA sends POST /api/v1/bookings with JWT → API validates slot availability",
        "API creates pending booking record → returns booking ID",
        "For online payment: SPA calls POST /api/v1/payments/order → API creates Razorpay order",
        "SPA opens Razorpay checkout modal → customer completes payment",
        "Razorpay callback → SPA sends POST /api/v1/payments/verify with signature",
        "API verifies signature server-side → updates booking to 'confirmed'",
        "Supabase broadcasts booking UPDATE → customer sees live status",
        "SPA navigates to confirmation screen (Step 4)",
    ])

    # 6. APPLICATION LAYER
    add_heading(doc, "6. Application Layer Architecture", 1)
    add_para(doc, "The application is organized into logical modules aligned with business domains, not UI screens.")
    add_table(doc, ["Module", "Responsibility", "Key API Domains"], [
        ["Discovery", "Salon search, listing, filters, featured/nearby", "/salons, /salons/search, /salons/nearby"],
        ["Salon Profile", "Detail, services, staff, gallery, reviews", "/salons/:id, /services, /staff, /reviews"],
        ["Booking", "Appointment creation, slots, cancel, reschedule", "/bookings, /salons/:id/slots"],
        ["Queue", "Walk-in join, advance, manual add, position tracking", "/salons/:id/queue/*"],
        ["Payment", "Razorpay order, verify, refund, history", "/payments/*"],
        ["Auth", "Login, register, OTP, Google, token refresh", "/auth/*"],
        ["User Profile", "Profile, avatar, booking history, wishlist", "/users/me, /wishlist"],
        ["Dashboard", "Salon owner KPIs, analytics, earnings", "/salons/:id/dashboard, /analytics/*"],
        ["Admin", "Platform management, salon approval, user moderation", "/admin/*"],
        ["Upload", "Image proxy to Cloudinary", "/upload/image"],
    ], [1.2, 2.5, 2.6])

    # 7. CLIENT ARCHITECTURE
    add_heading(doc, "7. Client Application Architecture", 1)
    add_para(doc,
        "The web client is a Single Page Application (SPA) built with React 18 and bundled by Vite. "
        "It communicates exclusively with the REST API for data mutations and reads, with Supabase "
        "used only for real-time subscriptions (queue and booking status).")

    add_heading(doc, "7.1 Client Layer Model", 2)
    add_table(doc, ["Layer", "Purpose", "Implementation"], [
        ["Routing Layer", "URL-based navigation, route guards, deep linking", "React Router v6 (createBrowserRouter)"],
        ["Page Layer", "Top-level views: Home, Salon List, Detail, Booking, Profile, Dashboard, Admin", "React page components"],
        ["Component Layer", "Reusable UI: cards, modals, forms, booking steps, dashboard widgets", "React functional components + CSS Modules"],
        ["State Layer", "Global app state, async thunks, domain slices", "Redux Toolkit (8 slices)"],
        ["Service Layer", "HTTP calls to API, token management, upload helpers", "Axios instance + domain service modules"],
        ["Real-Time Layer", "Live queue/booking subscriptions", "Supabase JS client"],
        ["Integration Layer", "Razorpay checkout, Google OAuth SDK", "Third-party SDKs loaded in browser"],
    ], [1.3, 2.5, 2.5])

    add_heading(doc, "7.2 Redux State Domains", 2)
    add_table(doc, ["Slice", "Domain State", "Async Operations"], [
        ["auth", "user, isAuthenticated, role, loading, error", "loginUser, registerUser, verifyOTP, googleLogin, logoutUser"],
        ["salons", "salon list, filters, search query, map data", "fetchSalons, fetchFeatured, searchSalons"],
        ["booking", "step, selectedService, date, slot, payment, currentBooking", "submitBooking, verifyPayment"],
        ["wishlist", "wishlisted salon IDs", "toggleWishlist (local + API sync)"],
        ["ui", "auth modal, notifications (toast), mobile menu", "openAuthModal, showNotification"],
        ["dashboard", "queue, bookings, earnings, activeView", "fetchDashboard, advanceQueue"],
        ["profile", "user profile, booking history, loyalty, settings", "fetchProfile, updateProfile"],
        ["admin", "platform stats, salons, users, activeView", "fetchAdminData, approveSalon"],
    ], [1.0, 2.5, 2.8])

    add_heading(doc, "7.3 Route Map", 2)
    add_table(doc, ["Route", "Access", "Purpose"], [
        ["/", "Public", "Marketing homepage, search, featured salons"],
        ["/salons", "Public", "Salon listing with filters and map"],
        ["/salons/:id", "Public", "Salon detail with inline booking panel"],
        ["/booking/:salonId", "Authenticated", "Multi-step booking flow (Steps 1–4)"],
        ["/profile", "Authenticated", "Customer profile, history, wishlist, settings"],
        ["/dashboard", "Shop Owner", "Salon operations dashboard"],
        ["/admin", "Admin", "Platform administration console"],
    ], [1.5, 1.2, 3.6])

    add_heading(doc, "7.4 Client Build & Runtime", 2)
    add_table(doc, ["Aspect", "Detail"], [
        ["Build tool", "Vite 5 with @vitejs/plugin-react"],
        ["Module system", "ES Modules (type: module)"],
        ["Path alias", "@ → src/ (configured in vite.config.js)"],
        ["CSS strategy", "CSS Modules per component + global.css design tokens"],
        ["UI libraries", "Bootstrap 5 (base), MUI 5 (select components), Framer Motion (animations)"],
        ["Output", "Static assets → dist/ folder for CDN/hosting deployment"],
        ["Env variables", "VITE_* prefixed (injected at build time)"],
    ], [1.5, 5.0])

    # 8. BACKEND API
    add_heading(doc, "8. Backend API Architecture", 1)
    add_para(doc,
        "The backend is a versioned REST API hosted at https://api.salonbazar.shop/api/v1. "
        "All endpoints return JSON. Authentication uses Bearer JWT tokens in the Authorization header.")

    add_heading(doc, "8.1 API Versioning & Base URL", 2)
    add_bullets(doc, [
        "Base URL: https://api.salonbazar.shop/api/v1",
        "Version prefix: /api/v1 (breaking changes require /api/v2)",
        "Content-Type: application/json (multipart/form-data for uploads)",
        "Timeout: 15 seconds (client-side axios default)",
    ])

    add_heading(doc, "8.2 Complete Endpoint Catalog", 2)

    endpoint_groups = [
        ("Authentication", [
            ("POST", "/auth/login", "Email + password login → JWT + refresh token"),
            ("POST", "/auth/register", "Register new user (customer or shop_owner)"),
            ("POST", "/auth/otp/send", "Send OTP to phone number"),
            ("POST", "/auth/otp/verify", "Verify OTP → JWT"),
            ("POST", "/auth/google", "Google ID token exchange → JWT"),
            ("POST", "/auth/logout", "Invalidate refresh token"),
            ("POST", "/auth/refresh", "Exchange refresh token for new access token"),
            ("POST", "/auth/forgot-password", "Send password reset email"),
            ("POST", "/auth/reset-password", "Reset password with token"),
        ]),
        ("Salons", [
            ("GET", "/salons", "List salons with filters (page, category, city, rating, price)"),
            ("GET", "/salons/featured", "Featured salons for homepage"),
            ("GET", "/salons/nearby", "Geo-based nearby salons (lat, lng, radius)"),
            ("GET", "/salons/:id", "Full salon detail"),
            ("GET", "/salons/search", "Full-text search (q + filters)"),
            ("POST", "/salons/register", "Salon owner registration"),
            ("PATCH", "/salons/:id", "Update salon profile"),
            ("DELETE", "/salons/:id", "Delete salon (admin)"),
            ("GET", "/salons/:id/qr", "Get QR code data for walk-in queue"),
        ]),
        ("Services & Staff", [
            ("GET", "/salons/:id/services", "List services for salon"),
            ("POST", "/salons/:id/services", "Create service"),
            ("PATCH", "/salons/:id/services/:sid", "Update service"),
            ("DELETE", "/salons/:id/services/:sid", "Delete service"),
            ("GET", "/salons/:id/staff", "List staff members"),
            ("POST/PATCH/DELETE", "/salons/:id/staff/:stid", "Staff CRUD"),
        ]),
        ("Bookings", [
            ("POST", "/bookings", "Create booking (online/qr/manual)"),
            ("GET", "/bookings", "List user's bookings"),
            ("GET", "/bookings/:id", "Booking detail"),
            ("PATCH", "/bookings/:id/cancel", "Cancel booking"),
            ("PATCH", "/bookings/:id/reschedule", "Reschedule booking"),
            ("GET", "/salons/:id/bookings", "Salon owner's bookings"),
            ("GET", "/salons/:id/slots", "Available time slots"),
            ("POST", "/bookings/backdated", "Add backdated completed entry"),
        ]),
        ("Queue", [
            ("POST", "/salons/:id/queue/join", "Customer joins walk-in queue"),
            ("GET", "/salons/:id/queue", "Full queue (owner view)"),
            ("GET", "/salons/:id/queue/me", "Customer's queue position"),
            ("DELETE", "/salons/:id/queue/leave", "Leave queue"),
            ("POST", "/salons/:id/queue/advance", "Owner advances queue"),
            ("POST", "/salons/:id/queue/manual", "Manual walk-in entry"),
        ]),
        ("Payments", [
            ("POST", "/payments/order", "Create Razorpay order"),
            ("POST", "/payments/verify", "Verify payment signature"),
            ("POST", "/payments/:id/refund", "Process refund"),
            ("GET", "/payments/history", "User payment history"),
        ]),
        ("Reviews & Wishlist", [
            ("GET/POST", "/salons/:id/reviews", "List/create reviews"),
            ("GET/POST/DELETE", "/wishlist", "Wishlist CRUD"),
        ]),
        ("User & Dashboard", [
            ("GET/PATCH", "/users/me", "Profile read/update"),
            ("POST", "/users/me/avatar", "Avatar upload"),
            ("GET", "/users/me/bookings", "Booking history"),
            ("GET", "/salons/:id/dashboard", "Salon dashboard overview"),
            ("GET", "/salons/:id/analytics/*", "Earnings, peak hours, services"),
            ("GET", "/admin/analytics", "Platform-wide admin analytics"),
        ]),
        ("Upload", [
            ("POST", "/upload/image", "Upload image → Cloudinary CDN URL"),
        ]),
    ]

    for group_name, endpoints in endpoint_groups:
        add_heading(doc, f"8.2.{group_name}", 3)
        add_table(doc, ["Method", "Endpoint", "Description"], endpoints, [0.8, 2.2, 3.3])

    add_heading(doc, "8.3 Standard API Response Format", 2)
    add_code_block(doc, """
// Success
{ "success": true, "data": { ... }, "message": "Optional message" }

// Error
{ "success": false, "error": "Human-readable message", "code": "BOOKING_SLOT_UNAVAILABLE", "status": 409 }

// Paginated list
{ "success": true, "data": [...], "meta": { "page": 1, "limit": 20, "total": 142 } }
    """)

    add_heading(doc, "8.4 HTTP Status Code Usage", 2)
    add_table(doc, ["Code", "Usage"], [
        ["200", "Successful GET, PATCH"],
        ["201", "Successful POST (resource created)"],
        ["204", "Successful DELETE"],
        ["400", "Validation error, bad request body"],
        ["401", "Missing or expired JWT"],
        ["403", "Insufficient role/permission"],
        ["404", "Resource not found"],
        ["409", "Conflict (slot taken, duplicate queue entry)"],
        ["422", "Business rule violation"],
        ["429", "Rate limit exceeded"],
        ["500", "Internal server error"],
    ], [0.8, 5.5])

    # 9. DATA ARCHITECTURE
    add_heading(doc, "9. Data Architecture", 1)
    add_heading(doc, "9.1 Primary Database — PostgreSQL", 2)
    add_para(doc, "PostgreSQL is the system of record for all transactional data. Supabase may host this Postgres instance and provide Realtime CDC.")

    add_heading(doc, "9.2 Entity-Relationship Overview", 2)
    add_table(doc, ["Entity", "Primary Key", "Key Fields", "Relationships"], [
        ["users", "id (UUID)", "name, email, phone, role, avatar_url", "1:N bookings, reviews, wishlist"],
        ["salons", "id (UUID)", "name, address, city, rating, status, plan", "1:N services, staff, bookings, queue"],
        ["services", "id (UUID)", "salon_id, name, category, duration, price", "N:1 salon; referenced by booking"],
        ["staff", "id (UUID)", "salon_id, name, role, rating, available", "N:1 salon; optional FK on booking"],
        ["bookings", "id (UUID)", "user_id, salon_id, service_id, date, time, status, total", "N:1 user, salon; 1:1 payment"],
        ["payments", "id (UUID)", "booking_id, method, amount, razorpay_ids, status", "N:1 booking"],
        ["queues", "id (UUID)", "salon_id, user_id, position, status, joined_at", "N:1 salon, user"],
        ["reviews", "id (UUID)", "salon_id, user_id, rating, text, verified", "N:1 salon, user"],
        ["wishlist", "id (UUID)", "user_id, salon_id, added_at", "N:1 user, salon"],
        ["coupons", "code (VARCHAR)", "discount, min_order, expiry, usage_limit", "Applied at booking"],
    ], [1.0, 1.0, 2.0, 2.3])

    add_heading(doc, "9.3 Booking Status State Machine", 2)
    add_code_block(doc, """
pending ──► confirmed ──► in_progress ──► completed
   │            │
   │            └──► cancelled (by user or salon, per policy)
   │
   └──► expired (payment timeout)
   
no_show ── (set by salon after confirmed + missed appointment)
    """)

    add_heading(doc, "9.4 Queue Status State Machine", 2)
    add_code_block(doc, """
waiting ──► in_progress ──► completed
   │
   ├──► left (customer leaves queue)
   └──► no_show (customer did not respond when called)
    """)

    add_heading(doc, "9.5 Indexing Strategy", 2)
    add_table(doc, ["Table", "Index", "Purpose"], [
        ["salons", "idx_salons_city_category", "Filter by city and category"],
        ["salons", "idx_salons_location (GiST)", "Geo proximity queries"],
        ["bookings", "idx_bookings_salon_date", "Dashboard daily appointments"],
        ["bookings", "idx_bookings_user_status", "Customer booking history"],
        ["queues", "idx_queues_salon_status", "Live queue queries"],
        ["reviews", "idx_reviews_salon_id", "Salon detail reviews"],
    ], [1.2, 2.0, 2.3])

    add_heading(doc, "9.6 Data Retention", 2)
    add_bullets(doc, [
        "Bookings and payments: retained indefinitely for audit and dispute resolution",
        "Queue entries: archived 30 days after completion",
        "Auth tokens: refresh tokens expire after 30 days; access tokens after 15 minutes",
        "Uploaded images: retained on Cloudinary; soft-delete with 90-day recovery window",
        "Audit logs (admin actions): retained 2 years",
    ])

    # 10. AUTH
    add_heading(doc, "10. Authentication & Authorization", 1)
    add_heading(doc, "10.1 Authentication Flow", 2)
    add_numbered(doc, [
        "User submits credentials (email/password, OTP, or Google token)",
        "API validates credentials → returns access JWT + refresh token",
        "Client stores tokens in localStorage (sb_token, sb_refresh_token)",
        "Axios request interceptor attaches Bearer token to all API calls",
        "On 401 response: client attempts POST /auth/refresh with refresh token",
        "If refresh succeeds: retry original request with new access token",
        "If refresh fails: clear tokens, redirect to /?auth=login",
    ])

    add_heading(doc, "10.2 JWT Token Structure", 2)
    add_code_block(doc, """
// Access Token Payload (example)
{
  "sub": "user-uuid",
  "email": "user@example.com",
  "role": "customer" | "shop_owner" | "admin",
  "salon_id": "salon-uuid",   // present for shop_owner
  "iat": 1710000000,
  "exp": 1710000900           // 15 min expiry
}
    """)

    add_heading(doc, "10.3 Role-Based Access Control (RBAC)", 2)
    add_table(doc, ["Role", "Code", "API Access Scope"], [
        ["Guest", "—", "Public endpoints only (salon list, detail, search)"],
        ["Customer", "customer", "Own bookings, profile, wishlist, queue join"],
        ["Shop Owner", "shop_owner", "Own salon data, dashboard, queue, bookings for owned salon"],
        ["Admin", "admin", "All platform data, salon approval, user management"],
    ], [1.2, 1.2, 3.9])

    add_heading(doc, "10.4 Authorization Enforcement", 2)
    add_bullets(doc, [
        "API middleware validates JWT on all protected routes",
        "Role checked per endpoint (e.g., /dashboard requires shop_owner or admin)",
        "Resource ownership verified (shop_owner can only access own salon_id)",
        "Client-side route guards provide UX redirect; server-side is authoritative",
    ])

    # 11. BOOKING
    add_heading(doc, "11. Booking Flow Architecture", 1)
    add_heading(doc, "11.1 Booking Entry Paths", 2)
    add_table(doc, ["Path", "Entry Point", "Initial Step", "Description"], [
        ["Full Flow", "/booking/:salonId", "Step 1 (Service)", "User completes all steps on booking page"],
        ["Express Flow", "Salon detail booking panel", "Step 3 (Payment)", "Pre-filled service/date/time from panel"],
        ["QR Walk-In", "QR scan at salon", "Queue join", "No pre-scheduled slot; joins live queue"],
        ["Manual Walk-In", "Salon dashboard", "Queue add", "Owner adds customer manually"],
    ], [1.2, 1.8, 1.3, 2.0])

    add_heading(doc, "11.2 Slot Availability Algorithm", 2)
    add_numbered(doc, [
        "Client requests GET /salons/:id/slots?service_id=X&date=YYYY-MM-DD",
        "API loads salon working hours for requested day of week",
        "API generates time slots based on service duration and salon slot interval (e.g., 30 min)",
        "API subtracts already-booked slots and staff unavailability",
        "API returns array of available slots with capacity count",
        "On booking POST: API uses database row lock (SELECT FOR UPDATE) to prevent double booking",
    ])

    add_heading(doc, "11.3 Booking State in Client", 2)
    add_para(doc, "The booking slice manages a 4-step wizard with the following state machine on the client:")
    add_table(doc, ["State Field", "Type", "Set At"], [
        ["step", "1 | 2 | 3 | 4", "initializeBookingFlow based on pre-filled data"],
        ["selectedService", "Service object", "Step 1 or salon detail panel"],
        ["selectedDate", "{ day, month, year, displayDate }", "Step 2 or booking panel"],
        ["selectedSlot", "string ('10:30 AM')", "Step 2 or booking panel"],
        ["selectedStaff", "Staff object | null", "Step 1 or booking panel (optional)"],
        ["paymentMethod", "'online' | 'counter'", "Step 3"],
        ["couponCode / couponDiscount", "string / number", "Step 3"],
        ["currentBooking", "Full booking object", "After submitBooking fulfilled"],
    ], [1.5, 2.0, 2.8])

    # 12. PAYMENT
    add_heading(doc, "12. Payment Architecture (Razorpay)", 1)
    add_heading(doc, "12.1 Payment Sequence Diagram", 2)
    add_numbered(doc, [
        "Client: POST /payments/order { booking_id, amount, currency: 'INR' }",
        "API: Creates Razorpay order server-side → returns { id, amount, currency }",
        "Client: Opens Razorpay checkout modal with order_id and VITE_RAZORPAY_KEY_ID",
        "Customer: Completes payment in Razorpay UI (UPI/card/netbanking)",
        "Razorpay: Returns razorpay_payment_id, razorpay_order_id, razorpay_signature to client handler",
        "Client: POST /payments/verify { razorpay_order_id, razorpay_payment_id, razorpay_signature, booking_id }",
        "API: Verifies HMAC signature using Razorpay secret → marks payment confirmed → booking confirmed",
        "On failure: Client shows error toast; booking remains in 'pending' state",
    ])

    add_heading(doc, "12.2 Pay-at-Counter Flow", 2)
    add_numbered(doc, [
        "Customer selects 'Pay at Salon' on Step 3",
        "Client: POST /bookings with payment_method: 'counter'",
        "API: Creates booking with status 'confirmed' and payment_status 'pending_at_counter'",
        "Salon marks payment received in dashboard upon customer arrival",
    ])

    add_heading(doc, "12.3 Security Requirements for Payments", 2)
    add_bullets(doc, [
        "Razorpay secret key stored ONLY on server — never in client env except public key_id",
        "Payment signature verification MUST happen server-side before confirming booking",
        "Amount in Razorpay order must match server-calculated total (prevent tampering)",
        "Webhook endpoint for Razorpay payment.captured events as backup confirmation",
        "Idempotency key on payment verify to prevent duplicate confirmations",
    ])

    # 13. QUEUE & REALTIME
    add_heading(doc, "13. Queue & Real-Time Architecture", 1)
    add_heading(doc, "13.1 Queue Data Flow", 2)
    add_numbered(doc, [
        "Customer scans salon QR → navigates to queue join page with salon_id",
        "POST /salons/:id/queue/join → API inserts row in queues table with next position",
        "Supabase Realtime broadcasts INSERT on queues table filtered by salon_id",
        "Salon dashboard subscription receives update → queue list re-renders",
        "Owner calls POST /salons/:id/queue/advance → current in_progress → completed; next waiting → in_progress",
        "Supabase broadcasts UPDATE → customer app updates position/turn notification",
    ])

    add_heading(doc, "13.2 Supabase Realtime Channels", 2)
    add_table(doc, ["Channel", "Table", "Filter", "Events", "Used By"], [
        ["queue:salon_{id}", "queues", "salon_id=eq.{id}", "INSERT, UPDATE, DELETE", "Dashboard queue view, customer tracker"],
        ["booking:{id}", "bookings", "id=eq.{id}", "UPDATE", "Customer booking status page"],
    ], [1.3, 0.8, 1.2, 1.2, 2.0])

    add_heading(doc, "13.3 Fallback When Realtime Unavailable", 2)
    add_bullets(doc, [
        "Client falls back to polling GET /salons/:id/queue/me every 15 seconds",
        "Dashboard polls GET /salons/:id/queue every 10 seconds",
        "Toast notification informs user if realtime connection drops",
    ])

    # 14. STATE MANAGEMENT
    add_heading(doc, "14. State Management Architecture", 1)
    add_para(doc, "Redux Toolkit is the single source of truth for client-side application state.")
    add_heading(doc, "14.1 Async Data Pattern", 2)
    add_code_block(doc, """
// Pattern: createAsyncThunk → extraReducers
export const submitBooking = createAsyncThunk('booking/submit', async (data) => {
  return await bookingService.create(data)  // Axios → API
})

// Slice handles pending/fulfilled/rejected → updates loading, error, currentBooking, step
    """)

    add_heading(doc, "14.2 Local vs Server State", 2)
    add_table(doc, ["State Type", "Storage", "Examples"], [
        ["Server state (cached)", "Redux slices", "Salon list, bookings, queue, profile"],
        ["UI state", "Redux ui slice + component useState", "Modal open, active tab, form inputs"],
        ["Auth tokens", "localStorage", "sb_token, sb_refresh_token"],
        ["Wishlist (optimistic)", "Redux + API sync", "Wishlist IDs"],
        ["Booking wizard", "Redux booking slice", "Persists across route navigation within session"],
    ], [1.5, 1.5, 3.3])

    add_heading(doc, "14.3 Notification System (Toast)", 2)
    add_para(doc, "User feedback uses a Redux-driven toast system — not browser alert() dialogs.")
    add_code_block(doc, """
dispatch(showNotification({ message: 'Please select a service first', type: 'warning' }))
// Types: 'success' | 'error' | 'warning'
// Rendered by global Notification component; auto-dismiss after 3.5 seconds
    """)

    # 15. SERVICE LAYER
    add_heading(doc, "15. Service Integration Layer", 1)
    add_para(doc, "The client service layer wraps the Axios API instance into domain-specific modules.")
    add_table(doc, ["Service Module", "Responsibility", "Endpoints Used"], [
        ["api.js", "Axios instance, JWT interceptors, token refresh, error formatting", "All"],
        ["authService.js", "Login, register, OTP, Google, token storage helpers", "/auth/*"],
        ["salonService.js", "Salon CRUD, search, featured, nearby, services", "/salons/*"],
        ["bookingService.js", "Create, cancel, reschedule, slots, backdated", "/bookings/*"],
        ["queueService.js", "Join, leave, advance, manual add, position", "/queue/*"],
        ["paymentService.js", "Create order, open Razorpay modal, verify, refund", "/payments/*"],
        ["uploadService.js", "Multipart image upload via API proxy", "/upload/image"],
        ["supabaseClient.js", "Realtime subscriptions for queue and bookings", "Supabase channels"],
    ], [1.3, 2.5, 2.5])

    add_heading(doc, "15.1 Axios Interceptor Architecture", 2)
    add_code_block(doc, """
Request Interceptor:
  → Read sb_token from localStorage
  → Attach Authorization: Bearer {token}

Response Interceptor:
  → Unwrap response.data automatically
  → On 401: attempt token refresh via /auth/refresh
  → On refresh failure: clear tokens, redirect /?auth=login
  → Format error as { message, status } for UI consumption
    """)

    # 16. MEDIA
    add_heading(doc, "16. Media & File Upload Architecture", 1)
    add_para(doc, "Images are never uploaded directly from browser to Cloudinary. The API proxies uploads to keep credentials secure.")
    add_numbered(doc, [
        "Client selects file → uploadService.uploadImage(file, folder)",
        "POST /upload/image with multipart/form-data { file, folder }",
        "API receives file → uploads to Cloudinary with server-side credentials",
        "API returns { url: 'https://res.cloudinary.com/...', publicId: '...' }",
        "Client stores CDN URL in salon profile, gallery, or avatar field",
    ])
    add_table(doc, ["Folder", "Usage", "Max Size"], [
        ["salons", "Salon gallery and cover images", "5 MB"],
        ["avatars", "User and staff profile photos", "2 MB"],
        ["services", "Service category images", "2 MB"],
        ["general", "Miscellaneous uploads", "5 MB"],
    ], [1.0, 3.0, 1.2])

    # 17. API STANDARDS
    add_heading(doc, "17. API Design Standards", 1)
    add_bullets(doc, [
        "RESTful resource naming: plural nouns (/salons, /bookings)",
        "Nested resources for ownership: /salons/:id/services, /salons/:id/queue",
        "PATCH for partial updates; PUT avoided unless full replacement",
        "Query params for filtering, sorting, pagination: ?page=1&limit=20&sort=rating",
        "Consistent error response schema with machine-readable error codes",
        "Idempotency-Key header on POST /bookings and POST /payments/verify",
        "Rate limiting: 100 req/min per IP (public), 300 req/min per user (authenticated)",
        "CORS: allow salonbazar.shop and staging.salonbazar.shop origins only",
    ])

    # 18. SECURITY
    add_heading(doc, "18. Security Architecture", 1)
    add_table(doc, ["Layer", "Control", "Implementation"], [
        ["Transport", "TLS 1.2+ encryption", "HTTPS on all domains; HSTS headers"],
        ["Authentication", "JWT access + refresh tokens", "Short-lived access (15 min), refresh rotation"],
        ["Authorization", "RBAC middleware", "Role + resource ownership checks on every protected route"],
        ["Input Validation", "Schema validation on all inputs", "Server-side validation (e.g., Zod/Joi)"],
        ["XSS Prevention", "Output encoding, CSP headers", "React auto-escaping; strict CSP on API responses"],
        ["CSRF", "Not applicable for JWT Bearer API", "No cookie-based auth"],
        ["Secrets Management", "Env vars / secret manager", "Razorpay secret, DB creds, Cloudinary secret in vault"],
        ["PII Protection", "Encryption at rest", "Phone/email encrypted in DB; masked in logs"],
        ["Payment PCI", "Delegated to Razorpay", "No card data touches SalonBazar servers"],
        ["Rate Limiting", "Per-IP and per-user limits", "Prevent brute force and abuse"],
    ], [1.2, 2.0, 3.1])

    add_heading(doc, "18.1 OWASP Top 10 Mitigations", 2)
    add_table(doc, ["Risk", "Mitigation"], [
        ["Injection", "Parameterized queries (ORM); input validation"],
        ["Broken Auth", "JWT with expiry; refresh rotation; secure token storage"],
        ["Sensitive Data Exposure", "HTTPS; encrypt PII; no secrets in client bundle"],
        ["XXE", "JSON-only API; no XML parsing"],
        ["Broken Access Control", "RBAC middleware; ownership checks"],
        ["Security Misconfiguration", "Hardened defaults; env-specific configs; no debug in prod"],
        ["XSS", "React escaping; CSP; sanitize user-generated content (reviews)"],
        ["Insecure Deserialization", "JSON schema validation on all payloads"],
        ["Known Vulnerabilities", "Dependabot/npm audit; regular dependency updates"],
        ["Insufficient Logging", "Structured audit logs for auth, payment, admin actions"],
    ], [2.0, 4.5])

    # 19. INFRASTRUCTURE
    add_heading(doc, "19. Infrastructure & Deployment", 1)
    add_heading(doc, "19.1 Environment Topology", 2)
    add_table(doc, ["Environment", "Web URL", "API URL", "Purpose"], [
        ["Production", "https://salonbazar.shop", "https://api.salonbazar.shop", "Live users"],
        ["Staging", "https://staging.salonbazar.shop", "https://staging-api.salonbazar.shop", "QA, UAT, demos"],
        ["Development", "http://localhost:5173", "http://localhost:3000", "Local development"],
    ], [1.2, 2.0, 2.3, 1.0])

    add_heading(doc, "19.2 Recommended Production Infrastructure", 2)
    add_table(doc, ["Component", "Recommended Service", "Notes"], [
        ["Web SPA hosting", "Vercel / Netlify / Cloudflare Pages", "Static dist/ deploy; CDN edge caching"],
        ["API server", "AWS ECS / Railway / Render / DigitalOcean App Platform", "Auto-scaling container"],
        ["Database", "Supabase Postgres / AWS RDS", "Managed Postgres with backups"],
        ["Realtime", "Supabase Realtime", "Included with Supabase Postgres"],
        ["Media CDN", "Cloudinary", "Image optimization and delivery"],
        ["DNS", "Cloudflare", "salonbazar.shop DNS + SSL + DDoS protection"],
        ["Secrets", "AWS Secrets Manager / Doppler", "Razorpay, DB, Cloudinary credentials"],
        ["Email", "SendGrid / AWS SES", "Transactional emails"],
        ["SMS", "MSG91 / Twilio", "OTP and booking reminders"],
    ], [1.3, 2.5, 2.5])

    add_heading(doc, "19.3 Deployment Architecture Diagram", 2)
    add_code_block(doc, """
[User Browser]
     │
     ▼
[Cloudflare CDN] ──► [Static SPA: salonbazar.shop]
     │
     ▼
[Load Balancer] ──► [API Server 1] [API Server 2] ... (auto-scaled)
                         │
              ┌──────────┼──────────┐
              ▼          ▼          ▼
         [PostgreSQL] [Redis*]  [Cloudinary]
              │
              ▼
         [Supabase Realtime]
         
* Redis optional: session cache, rate limiting, slot lock
    """)

    # 20. ENV CONFIG
    add_heading(doc, "20. Environment Configuration", 1)
    add_heading(doc, "20.1 Client Environment Variables (VITE_*)", 2)
    add_table(doc, ["Variable", "Example", "Purpose"], [
        ["VITE_APP_URL", "https://salonbazar.shop", "App base URL for redirects"],
        ["VITE_API_BASE_URL", "https://api.salonbazar.shop", "API server base URL"],
        ["VITE_SUPABASE_URL", "https://xxx.supabase.co", "Supabase project URL"],
        ["VITE_SUPABASE_ANON_KEY", "eyJ...", "Supabase anonymous key (public)"],
        ["VITE_RAZORPAY_KEY_ID", "rzp_live_xxx", "Razorpay public key (checkout modal)"],
        ["VITE_CLOUDINARY_CLOUD_NAME", "salonbazar", "Cloudinary cloud name (optional direct)"],
        ["VITE_GOOGLE_CLIENT_ID", "xxx.apps.googleusercontent.com", "Google OAuth client ID"],
    ], [2.0, 2.5, 2.0])

    add_heading(doc, "20.2 Server Environment Variables (Secret)", 2)
    add_table(doc, ["Variable", "Purpose"], [
        ["DATABASE_URL", "PostgreSQL connection string"],
        ["JWT_SECRET", "Access token signing key"],
        ["JWT_REFRESH_SECRET", "Refresh token signing key"],
        ["RAZORPAY_KEY_SECRET", "Razorpay server secret for signature verify"],
        ["CLOUDINARY_API_SECRET", "Cloudinary upload credentials"],
        ["SUPABASE_SERVICE_ROLE_KEY", "Server-side Supabase admin access"],
        ["GOOGLE_CLIENT_SECRET", "Google OAuth server secret"],
        ["SMS_API_KEY", "OTP SMS provider key"],
        ["EMAIL_API_KEY", "Transactional email provider key"],
    ], [2.5, 3.8])

    # 21. OBSERVABILITY
    add_heading(doc, "21. Observability & Monitoring", 1)
    add_table(doc, ["Concern", "Tool (Recommended)", "What to Monitor"], [
        ["Application errors", "Sentry", "JS errors (client), unhandled exceptions (API)"],
        ["API performance", "Datadog / New Relic", "Latency p50/p95/p99, error rate by endpoint"],
        ["Infrastructure", "CloudWatch / Grafana", "CPU, memory, disk, network on API servers"],
        ["Uptime", "UptimeRobot / Pingdom", "HTTP health checks on / and /api/v1/health"],
        ["Logs", "CloudWatch Logs / Loki", "Structured JSON logs with request_id correlation"],
        ["Real-user monitoring", "Sentry Performance / Web Vitals", "LCP, FID, CLS on client"],
        ["Payment monitoring", "Razorpay Dashboard + custom alerts", "Failed payments, refund requests"],
    ], [1.5, 2.0, 2.8])

    add_heading(doc, "21.1 Health Check Endpoint", 2)
    add_code_block(doc, """
GET /api/v1/health
Response 200:
{
  "status": "ok",
  "version": "1.0.0",
  "uptime": 86400,
  "database": "connected",
  "timestamp": "2026-06-07T10:00:00Z"
}
    """)

    # 22. PERFORMANCE
    add_heading(doc, "22. Performance & Scalability", 1)
    add_heading(doc, "22.1 Client Performance Targets", 2)
    add_table(doc, ["Metric", "Target", "Strategy"], [
        ["LCP (Largest Contentful Paint)", "< 2.5s", "CDN, image lazy loading, hero image optimization"],
        ["FID (First Input Delay)", "< 100ms", "Code splitting, defer non-critical JS"],
        ["CLS (Cumulative Layout Shift)", "< 0.1", "Skeleton loaders, explicit image dimensions"],
        ["Bundle size", "< 600KB gzipped", "Tree shaking, dynamic imports for admin/dashboard"],
        ["Time to Interactive", "< 3.5s on 4G", "Vite production build, preload critical assets"],
    ], [2.0, 1.2, 2.8])

    add_heading(doc, "22.2 API Performance Strategies", 2)
    add_bullets(doc, [
        "Database query optimization with proper indexes (see Section 9.5)",
        "Redis cache for salon list, featured salons (TTL 5 min)",
        "CDN cache for static salon gallery images via Cloudinary",
        "Pagination on all list endpoints (default limit: 20)",
        "Connection pooling on PostgreSQL (PgBouncer or ORM pool)",
        "Horizontal scaling of stateless API containers behind load balancer",
    ])

    add_heading(doc, "22.3 Scalability Projections", 2)
    add_table(doc, ["Scale", "Users", "Bookings/day", "Infrastructure"], [
        ["Launch", "5,000", "500", "1 API instance, managed Postgres (small)"],
        ["Growth", "50,000", "5,000", "2–3 API instances, Redis cache, read replica"],
        ["Scale", "500,000", "50,000", "Auto-scaling 5–10 instances, DB sharding by city"],
    ], [1.0, 1.0, 1.2, 3.3])

    # 23. ERROR HANDLING
    add_heading(doc, "23. Error Handling & Resilience", 1)
    add_table(doc, ["Layer", "Strategy", "User Experience"], [
        ["API", "Structured error codes + HTTP status", "Meaningful error messages in toast"],
        ["Axios interceptor", "Token refresh on 401; formatted reject", "Silent retry or login redirect"],
        ["Redux thunks", "rejected case sets error in slice", "Error displayed in form or toast"],
        ["React UI", "Error boundaries on page level", "Fallback UI instead of white screen"],
        ["Network", "Timeout 15s; retry button on failure", "Toast: 'Check your connection'"],
        ["Payment", "Razorpay failure handler → error toast", "Booking stays pending; user can retry"],
        ["Realtime", "Supabase disconnect → polling fallback", "Graceful degradation"],
    ], [1.0, 2.5, 2.8])

    # 24. CI/CD
    add_heading(doc, "24. CI/CD Pipeline", 1)
    add_heading(doc, "24.1 Recommended Pipeline Stages", 2)
    add_numbered(doc, [
        "Trigger: Push to main (production) or develop (staging); PR to main",
        "Lint: ESLint on src/ — fail on errors",
        "Build: npm run build (Vite) — fail on build errors",
        "Test: Unit tests (Vitest) + integration tests — fail on test failure",
        "Security: npm audit + SAST scan",
        "Deploy Staging: Auto-deploy dist/ to staging.salonbazar.shop on develop merge",
        "Deploy Production: Manual approval gate → deploy to salonbazar.shop on main merge",
        "Post-deploy: Smoke test health endpoint + critical user flows",
    ])

    add_heading(doc, "24.2 Branch Strategy", 2)
    add_table(doc, ["Branch", "Purpose", "Deploys To"], [
        ["main", "Production-ready code", "salonbazar.shop"],
        ["develop", "Integration branch", "staging.salonbazar.shop"],
        ["feature/*", "Feature development", "Preview deploy (optional)"],
        ["hotfix/*", "Emergency production fixes", "main (after review)"],
    ], [1.2, 2.5, 2.6])

    # 25. CURRENT VS TARGET
    add_heading(doc, "25. Current vs Target Architecture", 1)
    add_para(doc, "The web client prototype is substantially built. Backend integration is the primary gap.")
    add_table(doc, ["Component", "Current State", "Target State", "Gap"], [
        ["Web Client UI", "Complete with mock data", "Connected to live API", "Wire service calls to API"],
        ["Authentication", "Mock loginSuccess in modal", "JWT via authService + API", "Implement real auth flow"],
        ["Salon Data", "Mock constants + Redux", "API-driven with cache", "Replace mocks with API calls"],
        ["Booking Creation", "Mock submitBooking thunk", "POST /bookings + slot locking", "Backend + integration"],
        ["Payments", "Mock Razorpay flow", "Full Razorpay order/verify", "Backend payment endpoints"],
        ["Queue Realtime", "Mock dashboard data", "Supabase subscriptions live", "Connect supabaseClient"],
        ["Image Upload", "Service defined", "API → Cloudinary proxy live", "Backend upload endpoint"],
        ["Admin Panel", "Mock admin slice", "Live admin API", "Backend admin routes"],
        ["Notifications", "Toast system complete", "Push/SMS/email added", "Notification service"],
        ["Mobile Apps", "Not started", "React Native / Flutter", "Phase 2"],
    ], [1.3, 1.8, 1.8, 1.4])

    # 26. TECH STACK
    add_heading(doc, "26. Technology Stack Reference", 1)
    add_table(doc, ["Category", "Technology", "Version", "Purpose"], [
        ["Runtime (Client)", "React", "18.3", "UI framework"],
        ["Build", "Vite", "5.x", "Dev server and production bundler"],
        ["Language", "JavaScript (JSX)", "ES2022+", "Client application code"],
        ["Routing", "React Router", "6.23", "Client-side routing"],
        ["State", "Redux Toolkit", "2.2", "Global state management"],
        ["HTTP Client", "Axios", "1.15", "API communication"],
        ["CSS", "CSS Modules + Bootstrap 5", "5.3", "Component styling"],
        ["UI Components", "MUI 5, React Bootstrap", "5.15", "Select UI primitives"],
        ["Animation", "Framer Motion", "11.x", "Page and component animations"],
        ["Icons", "React Icons (Feather)", "5.2", "Icon library"],
        ["Carousel", "React Slick", "0.30", "Image carousels"],
        ["Realtime", "Supabase JS", "Latest", "Postgres CDC subscriptions"],
        ["Payments", "Razorpay Checkout", "SDK", "Payment modal"],
        ["Media CDN", "Cloudinary", "—", "Image storage and delivery"],
        ["Backend (Target)", "Node.js + Express / Fastify", "20 LTS", "REST API server"],
        ["Database (Target)", "PostgreSQL", "15+", "Primary data store"],
        ["Auth (Target)", "JWT + bcrypt", "—", "Token-based authentication"],
    ], [1.2, 2.0, 0.8, 2.3])

    # 27. ADRs
    add_heading(doc, "27. Architecture Decision Records (ADRs)", 1)

    adrs = [
        ("ADR-001: SPA over SSR", "Decision", "Use React SPA (Vite) instead of Next.js SSR for Phase 1",
         "Faster development velocity; SEO handled via static meta tags and future SSR for salon pages in Phase 2"),
        ("ADR-002: Redux Toolkit for State", "Decision", "Use Redux Toolkit over Context API or Zustand",
         "Complex multi-step booking flow, multiple user roles, and async thunks benefit from centralized predictable state"),
        ("ADR-003: REST over GraphQL", "Decision", "REST API instead of GraphQL",
         "Simpler for team; well-defined resources; easier caching; mobile apps can consume same REST API"),
        ("ADR-004: Supabase for Realtime Only", "Decision", "Use Supabase Realtime for queue; REST for everything else",
         "Avoids splitting business logic; Supabase used only where low-latency push is needed"),
        ("ADR-005: Server-Proxied Uploads", "Decision", "Upload images via API proxy, not direct Cloudinary from browser",
         "Keeps Cloudinary API secret off client; enables server-side validation and virus scanning"),
        ("ADR-006: JWT in localStorage", "Decision", "Store JWT in localStorage (not httpOnly cookies)",
         "SPA on separate domain from API; Bearer token pattern; mitigated by short expiry + refresh rotation"),
        ("ADR-007: Razorpay for Payments", "Decision", "Razorpay as sole payment gateway for India launch",
         "Best UPI coverage in India; PCI compliance delegated; well-documented SDK"),
        ("ADR-008: Mock-First Frontend", "Decision", "Build UI with mock data before API ready",
         "Parallel frontend/backend development; faster design iteration; API swap at integration phase"),
    ]
    for title, label, decision, rationale in adrs:
        add_heading(doc, title, 2)
        add_para(doc, f"{label}: {decision}")
        add_para(doc, f"Rationale: {rationale}")

    # 28. APPENDICES
    add_heading(doc, "28. Appendices", 1)

    add_heading(doc, "28.1 Appendix A — API Error Codes", 2)
    add_table(doc, ["Code", "HTTP", "Description"], [
        ["AUTH_INVALID_CREDENTIALS", "401", "Email/password incorrect"],
        ["AUTH_TOKEN_EXPIRED", "401", "JWT expired; refresh required"],
        ["AUTH_FORBIDDEN", "403", "Insufficient role for this action"],
        ["SALON_NOT_FOUND", "404", "Salon ID does not exist"],
        ["BOOKING_SLOT_UNAVAILABLE", "409", "Selected time slot no longer available"],
        ["BOOKING_CANCEL_TOO_LATE", "422", "Cancellation window passed (< 2 hours)"],
        ["PAYMENT_VERIFICATION_FAILED", "400", "Razorpay signature invalid"],
        ["QUEUE_ALREADY_JOINED", "409", "Customer already in queue at this salon"],
        ["COUPON_INVALID", "422", "Coupon code invalid or expired"],
        ["UPLOAD_FILE_TOO_LARGE", "413", "Image exceeds size limit"],
    ], [2.0, 0.7, 3.6])

    add_heading(doc, "28.2 Appendix B — Supabase Table Schema (queues)", 2)
    add_code_block(doc, """
CREATE TABLE queues (
  id          UUID PRIMARY KEY DEFAULT gen_random_uuid(),
  salon_id    UUID NOT NULL REFERENCES salons(id),
  user_id     UUID REFERENCES users(id),       -- nullable for walk-in without login
  customer_name VARCHAR(100),
  phone       VARCHAR(15),
  service_id  UUID REFERENCES services(id),
  position    INTEGER NOT NULL,
  status      VARCHAR(20) DEFAULT 'waiting',   -- waiting|in_progress|completed|left|no_show
  joined_at   TIMESTAMPTZ DEFAULT NOW(),
  updated_at  TIMESTAMPTZ DEFAULT NOW()
);
CREATE INDEX idx_queues_salon_status ON queues(salon_id, status);
    """)

    add_heading(doc, "28.3 Appendix C — Booking API Request/Response Example", 2)
    add_code_block(doc, """
POST /api/v1/bookings
Authorization: Bearer {jwt}

Request:
{
  "salon_id": "uuid",
  "service_id": "uuid",
  "staff_id": "uuid",           // optional
  "scheduled_date": "2026-06-15",
  "scheduled_time": "10:30 AM",
  "booking_type": "online",
  "payment_method": "online",
  "coupon_code": "FIRST50"
}

Response 201:
{
  "success": true,
  "data": {
    "id": "SB-A1B2C3",
    "status": "pending",
    "total": 468,
    "platform_fee": 19,
    "discount": 50,
    "salon_name": "Aura & Co.",
    "service_name": "Haircut & Styling",
    "display_date": "15 Jun 2026",
    "time": "10:30 AM"
  }
}
    """)

    add_heading(doc, "28.4 Appendix D — Related Documents", 2)
    add_bullets(doc, [
        "docs/PRD-SalonBazar.docx — Product Requirements Document",
        "docs/Technical-Architecture-SalonBazar.docx — This document",
        "API Specification (TBD) — OpenAPI/Swagger at api.salonbazar.shop/docs",
        "Database Schema Document (TBD)",
    ])

    add_heading(doc, "28.5 Document End", 2)
    add_para(doc, "— End of SalonBazar Technical Architecture Document —", bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"© {date.today().year} SalonBazar. Confidential — Internal Use.")
    run.font.size = Pt(10)
    run.font.italic = True

    doc.save(OUTPUT)
    print(f"Architecture document saved to {OUTPUT}")


if __name__ == "__main__":
    build_document()
