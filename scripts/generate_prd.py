"""Generate SalonBazar Product Requirements Document (PRD) as DOCX."""

from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "docs/PRD-SalonBazar.docx"


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1A1714")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_req_table(doc, requirements):
    add_table(
        doc,
        ["Req ID", "Module", "Requirement", "Priority", "Status"],
        requirements,
        col_widths=[0.9, 1.1, 3.2, 0.7, 0.9],
    )


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── COVER PAGE ──────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SalonBazar")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(26, 23, 20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Product Requirements Document (PRD)")
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(201, 168, 76)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Version: 1.0",
        f"Date: {date.today().strftime('%B %d, %Y')}",
        "Product: SalonBazar — Salon Discovery, Booking & Queue Platform",
        "Domain: https://salonbazar.shop",
        "API: https://api.salonbazar.shop/api/v1",
        "Prepared for: Product, Engineering, Design & Business Stakeholders",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # ── TABLE OF CONTENTS PLACEHOLDER ───────────────────────────
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Document Control",
        "2. Executive Summary",
        "3. Product Vision & Strategy",
        "4. Problem Statement & Opportunity",
        "5. Goals, Objectives & Success Metrics",
        "6. Target Market & User Personas",
        "7. Product Scope",
        "8. User Roles & Permissions",
        "9. End-to-End User Journeys",
        "10. Functional Requirements — Customer Platform",
        "11. Functional Requirements — Salon Owner Dashboard",
        "12. Functional Requirements — Platform Admin",
        "13. Booking & Payment Business Rules",
        "14. Queue & Walk-In System Requirements",
        "15. Notifications & Communication",
        "16. Search, Discovery & Personalization",
        "17. Reviews, Ratings & Trust",
        "18. Loyalty, Coupons & Monetization",
        "19. Non-Functional Requirements",
        "20. System Architecture Overview",
        "21. API & Integration Requirements",
        "22. Data Model & Entity Definitions",
        "23. Security, Privacy & Compliance",
        "24. Analytics & Reporting",
        "25. Release Plan & Roadmap",
        "26. Current Development Status",
        "27. Assumptions, Constraints & Dependencies",
        "28. Risks & Mitigation",
        "29. Open Questions",
        "30. User Stories (Detailed)",
        "31. Acceptance Criteria",
        "32. Error Handling & Edge Cases",
        "33. UI/UX Design Guidelines",
        "34. QA Test Scenarios",
        "35. Salon Partner Onboarding Checklist",
        "36. Glossary & Appendices",
    ]
    add_numbered(doc, toc_items)
    doc.add_page_break()

    # ── 1. DOCUMENT CONTROL ─────────────────────────────────────
    add_heading(doc, "1. Document Control", 1)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Document Title", "SalonBazar Product Requirements Document"],
            ["Version", "1.0"],
            ["Status", "Draft for Review"],
            ["Author", "Product Team"],
            ["Last Updated", date.today().strftime("%Y-%m-%d")],
            ["Confidentiality", "Internal / Stakeholder Use"],
            ["Distribution", "Product, Engineering, QA, Design, Business"],
        ],
        col_widths=[2.0, 4.5],
    )

    add_heading(doc, "1.1 Revision History", 2)
    add_table(
        doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["0.1", "2025-12-01", "Product Team", "Initial draft outline"],
            ["0.5", "2026-03-01", "Product Team", "Feature inventory from prototype"],
            ["1.0", date.today().strftime("%Y-%m-%d"), "Product Team", "Full PRD with requirements matrix"],
        ],
        col_widths=[0.7, 1.0, 1.3, 3.5],
    )

    add_heading(doc, "1.2 Approvals", 2)
    add_table(
        doc,
        ["Role", "Name", "Signature", "Date"],
        [
            ["Product Owner", "", "", ""],
            ["Engineering Lead", "", "", ""],
            ["Design Lead", "", "", ""],
            ["Business Sponsor", "", "", ""],
        ],
        col_widths=[1.5, 1.5, 2.0, 1.5],
    )

    add_heading(doc, "1.3 Related Documents", 2)
    add_bullets(
        doc,
        [
            "SalonBazar Brand & UX Guidelines (TBD)",
            "API Specification — api.salonbazar.shop (TBD)",
            "Database Schema Document (TBD)",
            "QA Test Plan (TBD)",
            "Salon Partner Onboarding Playbook (TBD)",
        ],
    )

    # ── 2. EXECUTIVE SUMMARY ────────────────────────────────────
    add_heading(doc, "2. Executive Summary", 1)
    add_para(
        doc,
        "SalonBazar is a digital marketplace and operations platform that connects customers with verified "
        "salons across India, starting with Ahmedabad. The platform enables customers to discover salons, "
        "book appointments online, join walk-in queues via QR codes, and manage their beauty service "
        "history — while giving salon owners a unified dashboard to manage bookings, live queues, "
        "revenue analytics, and customer communications.",
    )
    add_para(
        doc,
        "The product tagline — \"Your Best Look, Perfectly Booked\" — reflects the core value proposition: "
        "eliminate uncertainty, waiting, and friction from the salon experience. Customers gain transparency "
        "into availability, pricing, reviews, and real-time queue status. Salon partners gain tools to "
        "reduce no-shows, optimize staff utilization, and grow revenue through digital discovery.",
    )
    add_heading(doc, "2.1 Product at a Glance", 2)
    add_table(
        doc,
        ["Attribute", "Description"],
        [
            ["Product Name", "SalonBazar"],
            ["Product Type", "B2C Marketplace + B2B SaaS Operations Platform"],
            ["Primary Geography", "Ahmedabad, Gujarat (Phase 1); expandable pan-India"],
            ["Target Salons", "250+ verified salons (launch target)"],
            ["Customer Promise", "Discover, book, skip the queue, look amazing"],
            ["Partner Promise", "Grow bookings, manage queue, track revenue in real time"],
            ["Revenue Model", "Platform fee per booking, salon subscription tiers, featured listings"],
            ["Platforms", "Responsive Web (Phase 1); Mobile Apps (Phase 2)"],
        ],
        col_widths=[1.8, 4.7],
    )

    add_heading(doc, "2.2 Key Differentiators", 2)
    add_bullets(
        doc,
        [
            "Hybrid booking model: scheduled appointments AND walk-in queue via QR scan",
            "Live queue tracking with customer notifications when their turn approaches",
            "Verified salon listings with real photos, ratings, and transparent pricing",
            "Integrated payments (Razorpay) with pay-online or pay-at-counter options",
            "Salon owner dashboard with real-time KPIs, queue management, and analytics",
            "Platform admin console for salon onboarding, moderation, and platform health",
        ],
    )

    # ── 3. VISION & STRATEGY ────────────────────────────────────
    add_heading(doc, "3. Product Vision & Strategy", 1)
    add_heading(doc, "3.1 Vision Statement", 2)
    add_para(
        doc,
        "To become India's most trusted salon booking and queue management platform — where every customer "
        "walks in confident and every salon owner runs a smarter, more profitable business.",
        italic=True,
    )
    add_heading(doc, "3.2 Mission", 2)
    add_bullets(
        doc,
        [
            "Make salon services discoverable, bookable, and wait-free for urban consumers",
            "Digitize salon operations for SMB salon owners with minimal training overhead",
            "Build trust through verified listings, transparent pricing, and authentic reviews",
            "Create a scalable platform architecture supporting multi-city expansion",
        ],
    )
    add_heading(doc, "3.3 Strategic Pillars", 2)
    add_table(
        doc,
        ["Pillar", "Description", "Key Initiatives"],
        [
            ["Discovery", "Help users find the right salon fast", "Search, filters, categories, featured listings, map view"],
            ["Booking", "Frictionless appointment scheduling", "Multi-step booking, slot management, staff preference"],
            ["Queue", "Eliminate physical waiting", "QR walk-in, live queue, push/SMS notifications"],
            ["Operations", "Empower salon owners", "Dashboard, walk-in entry, analytics, earnings tracking"],
            ["Trust", "Build platform credibility", "Verified badges, reviews, ratings, cancellation policies"],
            ["Monetization", "Sustainable unit economics", "Platform fees, subscriptions, featured placement"],
        ],
        col_widths=[1.2, 2.3, 2.8],
    )

    # ── 4. PROBLEM STATEMENT ────────────────────────────────────
    add_heading(doc, "4. Problem Statement & Opportunity", 1)
    add_heading(doc, "4.1 Customer Pain Points", 2)
    add_bullets(
        doc,
        [
            "Unpredictable wait times at popular salons with no visibility into queue length",
            "Difficulty comparing salons by price, quality, availability, and location",
            "Phone-based booking is inconvenient; appointments are often lost or double-booked",
            "No centralized history of past services, preferred stylists, or loyalty benefits",
            "Limited payment options and unclear pricing before visiting the salon",
        ],
    )
    add_heading(doc, "4.2 Salon Owner Pain Points", 2)
    add_bullets(
        doc,
        [
            "Manual appointment books and WhatsApp messages create scheduling chaos",
            "Walk-in customers and pre-booked customers compete for the same staff capacity",
            "No-shows and last-minute cancellations reduce revenue with no recovery mechanism",
            "Limited digital presence beyond Instagram; discovery depends on word-of-mouth",
            "No analytics on peak hours, popular services, or staff performance",
        ],
    )
    add_heading(doc, "4.3 Market Opportunity", 2)
    add_para(
        doc,
        "India's beauty and wellness services market is large and fragmented. Ahmedabad alone has hundreds "
        "of independent salons with minimal digital tooling. SalonBazar targets the gap between generic "
        "marketplace listings and expensive enterprise salon software by offering a consumer-friendly "
        "marketplace paired with a lightweight, mobile-ready operations dashboard.",
    )

    # ── 5. GOALS & METRICS ──────────────────────────────────────
    add_heading(doc, "5. Goals, Objectives & Success Metrics", 1)
    add_heading(doc, "5.1 Business Goals (Year 1)", 2)
    add_table(
        doc,
        ["Goal", "Target", "Measurement"],
        [
            ["Salon network", "250+ verified salons in Ahmedabad", "Active salon count on platform"],
            ["Customer acquisition", "50,000+ registered users", "User registration count"],
            ["Booking volume", "10,000+ completed bookings/month", "Confirmed booking transactions"],
            ["Average rating", "4.5+ platform-wide", "Weighted average of salon ratings"],
            ["Salon retention", "80%+ monthly active salon partners", "Salons with ≥1 booking/month"],
            ["Revenue", "Platform fee + subscription ARR", "Monthly recurring revenue"],
        ],
        col_widths=[1.5, 2.2, 2.6],
    )

    add_heading(doc, "5.2 Product KPIs", 2)
    add_table(
        doc,
        ["KPI", "Definition", "Target"],
        [
            ["Booking conversion rate", "Bookings / salon detail page views", "≥ 8%"],
            ["Search-to-book time", "Median time from search to confirmed booking", "< 3 minutes"],
            ["No-show rate", "Missed appointments / total bookings", "< 10%"],
            ["Queue join rate", "QR walk-ins / total salon visits (tracked)", "≥ 30%"],
            ["Customer repeat rate", "Users with 2+ bookings in 90 days", "≥ 25%"],
            ["Payment success rate", "Successful Razorpay transactions / attempts", "≥ 95%"],
            ["NPS", "Net Promoter Score (quarterly survey)", "≥ 45"],
        ],
        col_widths=[1.6, 2.8, 1.1],
    )

    # ── 6. PERSONAS ─────────────────────────────────────────────
    add_heading(doc, "6. Target Market & User Personas", 1)
    personas = [
        ("Priya — Urban Professional (Customer)", "28, software engineer, Ahmedabad", [
            "Books hair/spa services 2–3 times per month",
            "Values convenience, ratings, and minimal wait time",
            "Prefers online payment and appointment reminders",
            "Uses mobile browser primarily; expects fast, modern UX",
        ]),
        ("Rahul — Grooming Enthusiast (Customer)", "32, business owner", [
            "Regular men's grooming customer",
            "Often walks in without appointment; frustrated by queues",
            "Would use QR queue if available at salon",
            "Price-sensitive but loyal to good stylists",
        ]),
        ("Meera — Salon Owner (Partner)", "42, owns premium unisex salon", [
            "Manages 5–8 staff members",
            "Currently uses WhatsApp + paper diary for bookings",
            "Wants more customers without hiring receptionist",
            "Needs revenue and peak-hour visibility",
        ]),
        ("Arjun — Platform Admin (Internal)", "Operations team member", [
            "Onboards and verifies new salon listings",
            "Moderates reviews and handles disputes",
            "Monitors platform health and analytics",
        ]),
    ]
    for name, demo, traits in personas:
        add_heading(doc, name, 2)
        add_para(doc, demo, italic=True)
        add_bullets(doc, traits)

    # ── 7. SCOPE ────────────────────────────────────────────────
    add_heading(doc, "7. Product Scope", 1)
    add_heading(doc, "7.1 In Scope (Phase 1 — MVP)", 2)
    add_bullets(
        doc,
        [
            "Public marketing homepage with search and category discovery",
            "Salon listing page with filters, sort, grid/list/map views",
            "Salon detail page with gallery, services, staff, reviews, location, inline booking",
            "Multi-step online booking flow (service → date/time → payment → confirmation)",
            "User authentication (email/password, OTP, Google OAuth)",
            "Customer profile with booking history, wishlist, settings",
            "Online payments via Razorpay and pay-at-counter option",
            "Coupon/discount application at checkout",
            "Salon owner dashboard (overview, queue, appointments, walk-in, analytics)",
            "Platform admin panel (salons, users, analytics, settings)",
            "Toast-based in-app notifications for validation and feedback",
            "QR-based walk-in queue (customer join + salon manage)",
        ],
    )
    add_heading(doc, "7.2 Out of Scope (Phase 1)", 2)
    add_bullets(
        doc,
        [
            "Native iOS/Android mobile applications",
            "In-salon POS and inventory management",
            "Multi-branch franchise management",
            "AI-based style recommendation engine",
            "Marketplace for beauty products/e-commerce",
            "International expansion and multi-currency support",
        ],
    )
    add_heading(doc, "7.3 Future Scope (Phase 2+)", 2)
    add_bullets(
        doc,
        [
            "Mobile apps with push notifications",
            "Salon subscription plans and billing portal",
            "Loyalty points redemption marketplace",
            "Staff scheduling and commission tracking",
            "WhatsApp booking bot integration",
            "Voice/IVR booking for non-smartphone users",
            "Multi-city expansion with localized SEO",
        ],
    )

    # ── 8. ROLES ────────────────────────────────────────────────
    add_heading(doc, "8. User Roles & Permissions", 1)
    add_table(
        doc,
        ["Role", "Description", "Key Permissions"],
        [
            ["Guest", "Unauthenticated visitor", "Browse salons, search, view details; must login to book"],
            ["Customer", "Registered end user", "Book, pay, review, wishlist, manage profile, join queue"],
            ["Shop Owner", "Salon business owner", "Manage salon profile, services, staff, queue, bookings, analytics"],
            ["Staff", "Salon employee (future)", "View assigned appointments, update service status"],
            ["Admin", "Platform operator", "Approve salons, manage users, platform settings, global analytics"],
        ],
        col_widths=[1.0, 1.8, 3.5],
    )

    add_heading(doc, "8.1 Permission Matrix (Summary)", 2)
    add_table(
        doc,
        ["Action", "Guest", "Customer", "Shop Owner", "Admin"],
        [
            ["View salon listings", "✓", "✓", "✓", "✓"],
            ["Book appointment", "—", "✓", "✓*", "✓"],
            ["Join walk-in queue", "—", "✓", "—", "—"],
            ["Manage live queue", "—", "—", "✓", "—"],
            ["Approve salon listing", "—", "—", "—", "✓"],
            ["Platform analytics", "—", "—", "Own salon", "✓"],
            ["Write review", "—", "✓", "—", "—"],
            ["Cancel booking", "—", "Own", "Own salon", "✓"],
        ],
        col_widths=[2.0, 0.7, 0.9, 1.0, 0.7],
    )
    add_para(doc, "* Shop owners may create walk-in/backdated bookings on behalf of customers.", italic=True, space_after=12)

    # ── 9. JOURNEYS ─────────────────────────────────────────────
    add_heading(doc, "9. End-to-End User Journeys", 1)

    journeys = [
        ("9.1 Customer — Discover & Book (Scheduled)", [
            "User lands on homepage → searches or browses categories",
            "User opens salon listing → applies filters (open now, category, price)",
            "User views salon detail → browses gallery, services, staff, reviews",
            "User selects service → picks date, time slot, optional staff preference",
            "User signs in (if not authenticated) → proceeds to booking checkout",
            "User selects payment method → applies coupon (optional) → confirms",
            "System creates booking → sends confirmation → shows success screen",
            "User can add to calendar, view booking in profile history",
        ]),
        ("9.2 Customer — Walk-In via QR", [
            "Customer arrives at salon → scans QR code displayed at reception",
            "Customer selects service (or confirms pre-selected) → joins queue",
            "System assigns queue position → customer tracks live status in app",
            "Customer receives notification when turn is approaching",
            "Salon staff advances queue → marks service in progress → completes",
            "Customer pays at counter or via app → optionally leaves review",
        ]),
        ("9.3 Salon Owner — Daily Operations", [
            "Owner logs into dashboard → views today's KPIs (bookings, revenue, queue)",
            "Owner monitors live queue → advances customers → adds walk-in manually",
            "Owner reviews upcoming appointments → handles cancellations/reschedules",
            "Owner checks analytics for peak hours and popular services",
            "Owner updates salon availability or staff schedule as needed",
        ]),
        ("9.4 Admin — Salon Onboarding", [
            "Salon owner submits registration via partner CTA",
            "Admin reviews application in admin panel → verifies documents",
            "Admin approves salon → salon goes live on marketplace",
            "Admin monitors platform metrics and handles user disputes",
        ]),
    ]
    for title, steps in journeys:
        add_heading(doc, title, 2)
        add_numbered(doc, steps)

    # ── 10. CUSTOMER REQUIREMENTS ───────────────────────────────
    add_heading(doc, "10. Functional Requirements — Customer Platform", 1)

    add_heading(doc, "10.1 Homepage & Discovery", 2)
    add_req_table(doc, [
        ["C-001", "Homepage", "Display hero section with search bar for salons, services, and areas", "P0", "UI Done"],
        ["C-002", "Homepage", "Show popular search chips (Hair Color, Bridal Makeup, etc.)", "P1", "UI Done"],
        ["C-003", "Homepage", "Display service category section for quick navigation", "P0", "UI Done"],
        ["C-004", "Homepage", "Show featured/top-rated salons near user location", "P0", "UI Done"],
        ["C-005", "Homepage", "Display nearby salons grid with filter pills (All, Open Now, Hair, Spa, etc.)", "P0", "UI Done"],
        ["C-006", "Homepage", "Show 'How It Works' 4-step explainer section", "P1", "UI Done"],
        ["C-007", "Homepage", "Show testimonials and partner CTA for salon owners", "P1", "UI Done"],
        ["C-008", "Homepage", "Search redirects to salon listing with query parameter", "P0", "UI Done"],
    ])

    add_heading(doc, "10.2 Salon Listing & Search", 2)
    add_req_table(doc, [
        ["C-010", "Listing", "Paginated/grid display of salons with image, rating, price, status", "P0", "UI Done"],
        ["C-011", "Listing", "Filter by category, price range, rating, distance, open now", "P0", "UI Done"],
        ["C-012", "Listing", "Sort by relevance, rating, distance, price", "P1", "UI Done"],
        ["C-013", "Listing", "Toggle grid view, list view, and map view", "P1", "UI Done"],
        ["C-014", "Listing", "Active filter chips with clear-all option", "P1", "UI Done"],
        ["C-015", "Listing", "Mobile filter sheet for small screens", "P1", "UI Done"],
        ["C-016", "Listing", "Wishlist toggle on salon cards", "P1", "UI Done"],
        ["C-017", "Listing", "Book Now navigates to salon detail for booking", "P0", "UI Done"],
        ["C-018", "Listing", "API integration for salon search and filters", "P0", "Pending"],
    ])

    add_heading(doc, "10.3 Salon Detail Page", 2)
    add_req_table(doc, [
        ["C-020", "Detail", "Hero with salon name, rating, location, open/closed status, share, wishlist", "P0", "UI Done"],
        ["C-021", "Detail", "Sticky section navigation (Gallery, Services, Staff, Reviews, Location)", "P1", "UI Done"],
        ["C-022", "Detail", "Photo gallery with lightbox/grid view", "P1", "UI Done"],
        ["C-023", "Detail", "About section with description, hours, amenities", "P1", "UI Done"],
        ["C-024", "Detail", "Services grouped by category with price, duration, popularity badge", "P0", "UI Done"],
        ["C-025", "Detail", "Staff profiles with photo, role, rating, availability", "P1", "UI Done"],
        ["C-026", "Detail", "Reviews with rating breakdown and verified badge", "P1", "UI Done"],
        ["C-027", "Detail", "Location section with address and map", "P1", "UI Done"],
        ["C-028", "Detail", "Desktop booking panel: service, calendar, time slots, staff preference", "P0", "UI Done"],
        ["C-029", "Detail", "Mobile bottom booking bar with expandable time slot picker", "P0", "UI Done"],
        ["C-030", "Detail", "Scrollable booking panel when content exceeds viewport", "P1", "UI Done"],
        ["C-031", "Detail", "Validation messages via toast (not browser alert)", "P1", "UI Done"],
        ["C-032", "Detail", "Walk-in QR option (scan to join queue)", "P1", "UI Partial"],
    ])

    add_heading(doc, "10.4 Booking Flow", 2)
    add_para(doc, "The booking flow supports two entry paths:", bold=True)
    add_bullets(doc, [
        "Full flow: User completes all steps on dedicated booking page (Steps 1–4)",
        "Express flow: User pre-selects service/date/time on salon page → lands on Step 3 (Payment)",
    ])
    add_req_table(doc, [
        ["C-040", "Booking", "Step 1: Select service from categorized list with staff preference", "P0", "UI Done"],
        ["C-041", "Booking", "Step 2: Select date from calendar and available time slot", "P0", "UI Done"],
        ["C-042", "Booking", "Step 3: Choose payment method (online Razorpay / pay at counter)", "P0", "UI Done"],
        ["C-043", "Booking", "Step 3: Apply and validate coupon codes", "P1", "UI Done"],
        ["C-044", "Booking", "Step 3: Display price breakdown (service, discount, platform fee, total)", "P0", "UI Done"],
        ["C-045", "Booking", "Step 3: Require cancellation policy agreement before confirm", "P0", "UI Done"],
        ["C-046", "Booking", "Step 4: Confirmation screen with booking ID, confetti, calendar export", "P0", "UI Done"],
        ["C-047", "Booking", "Sticky booking summary sidebar throughout flow", "P1", "UI Done"],
        ["C-048", "Booking", "Progress indicator showing current step (Service → Date → Payment → Done)", "P1", "UI Done"],
        ["C-049", "Booking", "Resume at correct step when arriving from salon detail with pre-filled data", "P0", "UI Done"],
        ["C-050", "Booking", "Auth guard: redirect unauthenticated users to login", "P0", "UI Done"],
        ["C-051", "Booking", "Backend booking creation API integration", "P0", "Pending"],
        ["C-052", "Booking", "Razorpay order creation and payment verification API", "P0", "Pending"],
    ])

    add_heading(doc, "10.5 Authentication & Profile", 2)
    add_req_table(doc, [
        ["C-060", "Auth", "Modal-based login and registration (email, password, phone)", "P0", "UI Done"],
        ["C-061", "Auth", "Google OAuth sign-in", "P1", "Planned"],
        ["C-062", "Auth", "OTP-based phone login", "P1", "Planned"],
        ["C-063", "Auth", "Forgot/reset password flow", "P1", "Planned"],
        ["C-064", "Auth", "JWT token storage with refresh token rotation", "P0", "Planned"],
        ["C-065", "Profile", "Profile hero with avatar, loyalty points, spend, visit count", "P1", "UI Done"],
        ["C-066", "Profile", "Booking history with status, cancel, rebook, review actions", "P0", "UI Done"],
        ["C-067", "Profile", "Wishlist tab with remove and book actions", "P1", "UI Done"],
        ["C-068", "Profile", "Settings: personal info, password, notification preferences", "P1", "UI Done"],
        ["C-069", "Profile", "Avatar upload via Cloudinary", "P2", "Planned"],
    ])

    # ── 11. SALON DASHBOARD ─────────────────────────────────────
    add_heading(doc, "11. Functional Requirements — Salon Owner Dashboard", 1)
    add_req_table(doc, [
        ["S-001", "Dashboard", "Auth-protected access for shop_owner role", "P0", "UI Done"],
        ["S-002", "Overview", "KPI cards: today's bookings, revenue, queue count, trends", "P0", "UI Done"],
        ["S-003", "Overview", "Today's appointment list with status indicators", "P0", "UI Done"],
        ["S-004", "Overview", "Quick action: Add Walk-In", "P1", "UI Done"],
        ["S-005", "Queue", "Live queue list with waiting/in-progress/completed states", "P0", "UI Done"],
        ["S-006", "Queue", "Advance queue (call next customer)", "P0", "UI Done"],
        ["S-007", "Queue", "Real-time badge count for waiting customers", "P1", "UI Done"],
        ["S-008", "Bookings", "Appointment calendar/list view with filters", "P0", "UI Done"],
        ["S-009", "Bookings", "Handle pending, confirmed, completed, cancelled statuses", "P0", "UI Done"],
        ["S-010", "Walk-In", "Manual walk-in entry form (name, phone, service)", "P1", "UI Done"],
        ["S-011", "Analytics", "Earnings chart, peak hours, service popularity", "P1", "UI Done"],
        ["S-012", "Settings", "Salon profile, services, staff, hours management", "P1", "UI Partial"],
        ["S-013", "QR", "Generate and display salon QR code for walk-in queue", "P1", "Planned"],
        ["S-014", "API", "Integrate dashboard endpoints with live data", "P0", "Pending"],
    ])

    # ── 12. ADMIN ───────────────────────────────────────────────
    add_heading(doc, "12. Functional Requirements — Platform Admin", 1)
    add_req_table(doc, [
        ["A-001", "Admin", "Auth-protected admin panel access", "P0", "UI Done"],
        ["A-002", "Overview", "Platform-wide KPIs: total salons, users, bookings, revenue", "P0", "UI Done"],
        ["A-003", "Salons", "List all salons with status (active, pending, suspended)", "P0", "UI Done"],
        ["A-004", "Salons", "Approve/reject pending salon registrations", "P0", "UI Done"],
        ["A-005", "Salons", "Update salon subscription plan", "P1", "UI Done"],
        ["A-006", "Users", "User list with role, status, activity", "P1", "UI Done"],
        ["A-007", "Users", "Suspend/activate user accounts", "P1", "UI Done"],
        ["A-008", "Analytics", "Platform analytics dashboard", "P1", "UI Done"],
        ["A-009", "Settings", "Platform configuration (fees, policies, feature flags)", "P2", "UI Done"],
        ["A-010", "API", "Admin API integration for all management actions", "P0", "Pending"],
    ])

    # ── 13. BUSINESS RULES ──────────────────────────────────────
    add_heading(doc, "13. Booking & Payment Business Rules", 1)
    add_table(
        doc,
        ["Rule ID", "Rule", "Details"],
        [
            ["BR-001", "Slot availability", "Time slots marked unavailable when fully booked or staff unavailable"],
            ["BR-002", "Past dates", "Users cannot select past calendar dates for booking"],
            ["BR-003", "Auth required", "Booking requires authenticated customer account"],
            ["BR-004", "Platform fee", "₹19 platform fee applied on online bookings (configurable by admin)"],
            ["BR-005", "Coupons", "Valid coupons: FIRST50 (₹50 off), SALON20 (20% off), BAZAR100 (₹100 off, min ₹500)"],
            ["BR-006", "Payment methods", "Online (Razorpay: UPI, cards, net banking) or Pay at Counter"],
            ["BR-007", "Cancellation", "Free cancellation up to 2 hours before appointment"],
            ["BR-008", "Confirmation", "Booking confirmed only after successful payment OR counter payment flag"],
            ["BR-009", "Staff preference", "Optional; system assigns best available if not specified"],
            ["BR-010", "Express booking", "Pre-filled salon page booking skips to payment step when all fields set"],
        ],
        col_widths=[0.8, 1.5, 3.9],
    )

    # ── 14. QUEUE ───────────────────────────────────────────────
    add_heading(doc, "14. Queue & Walk-In System Requirements", 1)
    add_req_table(doc, [
        ["Q-001", "Queue", "Customer joins queue via QR scan at salon", "P0", "Planned"],
        ["Q-002", "Queue", "Display estimated wait time based on queue length and avg service duration", "P1", "UI Partial"],
        ["Q-003", "Queue", "Customer sees live queue position in app", "P0", "Planned"],
        ["Q-004", "Queue", "Notification when customer is next in queue", "P1", "Planned"],
        ["Q-005", "Queue", "Salon owner advances queue from dashboard", "P0", "UI Done"],
        ["Q-006", "Queue", "Manual walk-in entry by salon staff", "P1", "UI Done"],
        ["Q-007", "Queue", "Queue states: waiting → in_progress → completed → no_show", "P0", "UI Done"],
        ["Q-008", "Queue", "Prevent duplicate queue entries for same customer at same salon", "P1", "Planned"],
    ])

    # ── 15. NOTIFICATIONS ───────────────────────────────────────
    add_heading(doc, "15. Notifications & Communication", 1)
    add_table(
        doc,
        ["Trigger", "Channel", "Recipient", "Message (Sample)"],
        [
            ["Booking confirmed", "Email + In-app", "Customer", "Your appointment at {salon} is confirmed for {date} {time}"],
            ["Booking reminder", "SMS + Push", "Customer", "Reminder: Your {service} appointment is in 2 hours"],
            ["Queue position update", "Push + In-app", "Customer", "You're #{n} in queue. ~{wait} min wait"],
            ["Turn approaching", "Push + SMS", "Customer", "You're up next at {salon}! Please proceed to reception"],
            ["Booking cancelled", "Email + In-app", "Customer + Salon", "Booking {id} has been cancelled"],
            ["Salon approved", "Email", "Shop Owner", "Congratulations! {salon} is now live on SalonBazar"],
            ["Validation error", "In-app toast", "User", "Please select a service first"],
            ["Payment failed", "In-app toast", "Customer", "Payment failed. Please try again."],
        ],
        col_widths=[1.4, 1.2, 1.1, 2.6],
    )

    # ── 16. SEARCH ──────────────────────────────────────────────
    add_heading(doc, "16. Search, Discovery & Personalization", 1)
    add_bullets(doc, [
        "Full-text search across salon names, services, areas, and tags",
        "Geolocation-based 'nearby' sorting using device GPS (with permission)",
        "Featured salons: minimum 4.8 rating and 100+ reviews required",
        "Category-based browsing: Hair, Skin, Bridal, Nails, Spa, Men's Grooming",
        "Personalized recommendations based on booking history (Phase 2)",
        "Recently viewed salons (Phase 2)",
    ])

    # ── 17. REVIEWS ─────────────────────────────────────────────
    add_heading(doc, "17. Reviews, Ratings & Trust", 1)
    add_req_table(doc, [
        ["R-001", "Reviews", "Display star rating and review count on all salon cards", "P0", "UI Done"],
        ["R-002", "Reviews", "Rating breakdown (5-star distribution) on detail page", "P1", "UI Done"],
        ["R-003", "Reviews", "Verified visit badge on reviews from confirmed bookings", "P1", "UI Done"],
        ["R-004", "Reviews", "Customer can submit review after completed booking", "P1", "UI Partial"],
        ["R-005", "Reviews", "Admin moderation for inappropriate reviews", "P2", "Planned"],
        ["R-006", "Trust", "Open/Closed live status based on working hours", "P0", "UI Done"],
        ["R-007", "Trust", "Verified salon badge for admin-approved listings", "P1", "Planned"],
    ])

    # ── 18. MONETIZATION ────────────────────────────────────────
    add_heading(doc, "18. Loyalty, Coupons & Monetization", 1)
    add_heading(doc, "18.1 Revenue Streams", 2)
    add_table(
        doc,
        ["Stream", "Description", "Phase"],
        [
            ["Platform booking fee", "Fixed fee per online booking (₹19 default)", "Phase 1"],
            ["Salon subscription", "Monthly plans: Free, Pro, Enterprise with feature tiers", "Phase 2"],
            ["Featured listings", "Paid placement on homepage and category pages", "Phase 2"],
            ["Promoted search", "Sponsored salon results in search", "Phase 3"],
        ],
        col_widths=[1.5, 3.0, 1.0],
    )
    add_heading(doc, "18.2 Loyalty Program (Phase 2)", 2)
    add_bullets(doc, [
        "Earn points on every completed booking (1 point per ₹10 spent)",
        "Redeem points for discounts on future bookings",
        "Tier levels: Silver, Gold, Platinum with increasing benefits",
        "Birthday and anniversary bonus offers",
    ])

    # ── 19. NFR ─────────────────────────────────────────────────
    add_heading(doc, "19. Non-Functional Requirements", 1)
    add_table(
        doc,
        ["Category", "Requirement", "Target"],
        [
            ["Performance", "Page load time (LCP)", "< 2.5 seconds on 4G"],
            ["Performance", "API response time (p95)", "< 500ms"],
            ["Performance", "Search results", "< 1 second"],
            ["Availability", "Platform uptime", "99.5% monthly"],
            ["Scalability", "Concurrent users", "10,000+ without degradation"],
            ["Scalability", "Bookings per minute", "500+ peak capacity"],
            ["Security", "Data encryption", "TLS 1.2+ in transit; AES-256 at rest"],
            ["Security", "Authentication", "JWT with refresh token rotation"],
            ["Accessibility", "WCAG compliance", "Level AA target"],
            ["Responsive", "Device support", "Mobile-first; 320px to 1920px+"],
            ["Browser", "Support", "Chrome, Safari, Firefox, Edge (last 2 versions)"],
            ["Localization", "Language", "English (Phase 1); Gujarati/Hindi (Phase 2)"],
        ],
        col_widths=[1.3, 2.5, 2.5],
    )

    # ── 20. ARCHITECTURE ────────────────────────────────────────
    add_heading(doc, "20. System Architecture Overview", 1)
    add_para(doc, "SalonBazar follows a modern three-tier architecture:", bold=True)
    add_bullets(doc, [
        "Presentation Layer: Responsive web application (customer, partner, admin interfaces)",
        "Application Layer: RESTful API server (api.salonbazar.shop/api/v1)",
        "Data Layer: Primary database + Supabase for real-time features + Cloudinary for media",
    ])
    add_heading(doc, "20.1 High-Level Component Diagram", 2)
    add_para(doc, "[ Customer Web App ] ←→ [ API Gateway / REST API ] ←→ [ Database ]")
    add_para(doc, "                              ↕                    ↕")
    add_para(doc, "                    [ Razorpay Payments ]   [ Supabase Realtime ]")
    add_para(doc, "                              ↕")
    add_para(doc, "                    [ Cloudinary CDN (Images) ]")
    add_para(doc, "                              ↕")
    add_para(doc, "                    [ Google OAuth / SMS OTP Provider ]")

    add_heading(doc, "20.2 Environments", 2)
    add_table(
        doc,
        ["Environment", "URL", "Purpose"],
        [
            ["Production", "salonbazar.shop / api.salonbazar.shop", "Live users"],
            ["Staging", "staging.salonbazar.shop", "QA and UAT"],
            ["Development", "localhost", "Local development"],
        ],
        col_widths=[1.2, 2.8, 2.5],
    )

    # ── 21. API ─────────────────────────────────────────────────
    add_heading(doc, "21. API & Integration Requirements", 1)
    add_heading(doc, "21.1 Authentication Endpoints", 2)
    add_table(
        doc,
        ["Method", "Endpoint", "Description"],
        [
            ["POST", "/api/v1/auth/login", "Email/password login"],
            ["POST", "/api/v1/auth/register", "New user registration"],
            ["POST", "/api/v1/auth/otp/send", "Send OTP to phone"],
            ["POST", "/api/v1/auth/otp/verify", "Verify OTP and authenticate"],
            ["POST", "/api/v1/auth/google", "Google OAuth token exchange"],
            ["POST", "/api/v1/auth/logout", "Invalidate session"],
            ["POST", "/api/v1/auth/refresh", "Refresh JWT token"],
            ["POST", "/api/v1/auth/forgot-password", "Initiate password reset"],
            ["POST", "/api/v1/auth/reset-password", "Complete password reset"],
        ],
        col_widths=[0.8, 2.5, 2.9],
    )

    add_heading(doc, "21.2 Core Business Endpoints", 2)
    add_table(
        doc,
        ["Domain", "Key Endpoints", "Purpose"],
        [
            ["Salons", "GET /salons, /salons/featured, /salons/nearby, /salons/:id", "Discovery & detail"],
            ["Services", "GET/POST/PUT/DELETE /salons/:id/services", "Service catalog CRUD"],
            ["Bookings", "POST /bookings, GET /bookings/:id, POST /bookings/:id/cancel", "Booking lifecycle"],
            ["Slots", "GET /salons/:id/slots", "Available time slots"],
            ["Queue", "POST /salons/:id/queue/join, GET /queue, POST /queue/advance", "Walk-in queue"],
            ["Reviews", "GET/POST /salons/:id/reviews", "Ratings & reviews"],
            ["Wishlist", "GET/POST/DELETE /wishlist", "Customer favorites"],
            ["Payments", "POST /payments/order, POST /payments/verify", "Razorpay integration"],
            ["Dashboard", "GET /salons/:id/dashboard, /analytics/*", "Salon owner analytics"],
            ["Admin", "GET /admin/analytics", "Platform-wide metrics"],
            ["Upload", "POST /upload/image", "Image upload to Cloudinary"],
        ],
        col_widths=[1.0, 3.2, 2.1],
    )

    add_heading(doc, "21.3 Third-Party Integrations", 2)
    add_table(
        doc,
        ["Service", "Purpose", "Status"],
        [
            ["Razorpay", "Online payment processing (UPI, cards, net banking)", "Configured; integration pending"],
            ["Supabase", "Real-time queue updates, optional auth/storage", "Configured; integration pending"],
            ["Cloudinary", "Salon gallery, staff photos, user avatars", "Configured; integration pending"],
            ["Google OAuth", "Social sign-in", "Configured; integration pending"],
            ["SMS Gateway", "OTP and booking reminders", "Planned"],
            ["Email Service", "Transactional emails (SendGrid/SES)", "Planned"],
        ],
        col_widths=[1.3, 3.0, 1.9],
    )

    # ── 22. DATA MODEL ──────────────────────────────────────────
    add_heading(doc, "22. Data Model & Entity Definitions", 1)
    add_table(
        doc,
        ["Entity", "Key Attributes", "Relationships"],
        [
            ["User", "id, name, email, phone, role, avatar, created_at", "Has many Bookings, Reviews, WishlistItems"],
            ["Salon", "id, name, category, address, city, rating, is_open, status, plan", "Has many Services, Staff, Bookings, Reviews"],
            ["Service", "id, salon_id, name, category, duration, price, popular, description", "Belongs to Salon; referenced in Booking"],
            ["Staff", "id, salon_id, name, role, rating, specialties, available", "Belongs to Salon; optional in Booking"],
            ["Booking", "id, user_id, salon_id, service_id, staff_id, date, time, status, total", "Belongs to User, Salon; has Payment"],
            ["Payment", "id, booking_id, method, amount, razorpay_ids, status", "Belongs to Booking"],
            ["QueueEntry", "id, salon_id, user_id, service, position, status, joined_at", "Belongs to Salon, User"],
            ["Review", "id, salon_id, user_id, rating, text, service, verified", "Belongs to Salon, User"],
            ["Coupon", "code, discount, min_order, expiry, usage_limit", "Applied to Booking"],
            ["WishlistItem", "user_id, salon_id, added_at", "Belongs to User, Salon"],
        ],
        col_widths=[1.0, 2.8, 2.5],
    )

    add_heading(doc, "22.1 Booking Status Lifecycle", 2)
    add_para(doc, "pending → confirmed → in_progress → completed | cancelled | no_show")

    # ── 23. SECURITY ────────────────────────────────────────────
    add_heading(doc, "23. Security, Privacy & Compliance", 1)
    add_bullets(doc, [
        "All API communication over HTTPS (TLS 1.2+)",
        "JWT access tokens with short expiry; refresh tokens stored securely",
        "Role-based access control (RBAC) enforced at API layer",
        "PII encryption at rest for phone numbers and email addresses",
        "PCI-DSS compliance delegated to Razorpay for payment card data",
        "GDPR-inspired data export and account deletion on user request",
        "Rate limiting on auth endpoints to prevent brute force",
        "Input validation and sanitization on all user-submitted content",
        "Audit logging for admin actions (salon approval, user suspension)",
    ])

    # ── 24. ANALYTICS ───────────────────────────────────────────
    add_heading(doc, "24. Analytics & Reporting", 1)
    add_heading(doc, "24.1 Salon Owner Analytics", 2)
    add_bullets(doc, [
        "Daily/weekly/monthly revenue charts",
        "Peak hours heatmap by day of week",
        "Top services by booking count and revenue",
        "Customer return rate and new vs returning split",
        "Average wait time and queue throughput",
        "Staff utilization and performance metrics",
    ])
    add_heading(doc, "24.2 Platform Admin Analytics", 2)
    add_bullets(doc, [
        "Total salons, users, bookings, GMV trends",
        "Salon onboarding funnel (registered → pending → active)",
        "Geographic distribution of salons and bookings",
        "Payment success/failure rates",
        "Platform fee revenue tracking",
    ])

    # ── 25. ROADMAP ─────────────────────────────────────────────
    add_heading(doc, "25. Release Plan & Roadmap", 1)
    add_table(
        doc,
        ["Phase", "Timeline", "Deliverables"],
        [
            ["Phase 1 — MVP", "Q1–Q2 2026", "Web app, booking flow, salon dashboard, admin panel, Razorpay, Ahmedabad launch"],
            ["Phase 2 — Growth", "Q3 2026", "Mobile apps, push notifications, loyalty program, salon subscriptions"],
            ["Phase 3 — Scale", "Q4 2026+", "Multi-city expansion, AI recommendations, WhatsApp bot, franchise support"],
        ],
        col_widths=[1.3, 1.2, 4.0],
    )

    add_heading(doc, "25.1 MVP Feature Priority", 2)
    add_table(
        doc,
        ["Priority", "Features"],
        [
            ["P0 — Must Have", "Salon discovery, detail, booking, payment, auth, profile, salon dashboard queue, admin approval"],
            ["P1 — Should Have", "Map view, coupons, analytics, walk-in QR, reviews submission, notification preferences"],
            ["P2 — Nice to Have", "Loyalty points, staff app, subscription billing, promoted listings"],
        ],
        col_widths=[1.5, 5.0],
    )

    # ── 26. DEV STATUS ──────────────────────────────────────────
    add_heading(doc, "26. Current Development Status", 1)
    add_para(
        doc,
        "The customer-facing web application prototype is substantially complete with mock data. "
        "Backend API integration is the primary remaining engineering work. The table below summarizes "
        "implementation status as of this document date.",
    )
    add_table(
        doc,
        ["Module", "UI/UX", "Frontend Logic", "Backend API", "Overall"],
        [
            ["Homepage & Discovery", "Complete", "Complete", "Pending", "70%"],
            ["Salon Listing & Filters", "Complete", "Complete", "Pending", "70%"],
            ["Salon Detail & Booking Panel", "Complete", "Complete", "Pending", "75%"],
            ["Multi-Step Booking Flow", "Complete", "Complete", "Pending", "70%"],
            ["Authentication", "Complete", "Mock auth", "Pending", "40%"],
            ["Customer Profile", "Complete", "Mock data", "Pending", "60%"],
            ["Payments (Razorpay)", "Complete", "Mock flow", "Pending", "45%"],
            ["Salon Owner Dashboard", "Complete", "Mock data", "Pending", "60%"],
            ["Platform Admin", "Complete", "Mock data", "Pending", "55%"],
            ["Queue & Walk-In", "Partial UI", "Partial", "Pending", "30%"],
            ["Notifications (Toast)", "Complete", "Complete", "N/A", "100%"],
            ["Reviews System", "Display done", "Partial", "Pending", "40%"],
        ],
        col_widths=[1.8, 1.0, 1.2, 1.0, 0.8],
    )

    # ── 27. ASSUMPTIONS ─────────────────────────────────────────
    add_heading(doc, "27. Assumptions, Constraints & Dependencies", 1)
    add_heading(doc, "27.1 Assumptions", 2)
    add_bullets(doc, [
        "Salon owners have smartphones and basic digital literacy",
        "Customers in target market have access to smartphones with internet",
        "Razorpay merchant account will be approved before payment launch",
        "Initial launch limited to Ahmedabad metropolitan area",
        "Salons will display QR codes at reception for walk-in queue",
    ])
    add_heading(doc, "27.2 Constraints", 2)
    add_bullets(doc, [
        "Phase 1 is web-only (no native mobile apps)",
        "Real-time queue requires stable internet at salon premises",
        "Platform fee structure subject to business team final approval",
        "Salon service catalog managed manually by salon owners initially",
    ])
    add_heading(doc, "27.3 Dependencies", 2)
    add_bullets(doc, [
        "Backend API development and deployment on api.salonbazar.shop",
        "Razorpay payment gateway integration and webhook setup",
        "Supabase project for real-time queue subscriptions",
        "Cloudinary account for image storage and CDN delivery",
        "SMS/Email provider for OTP and transactional notifications",
        "Domain and SSL certificates for production deployment",
    ])

    # ── 28. RISKS ───────────────────────────────────────────────
    add_heading(doc, "28. Risks & Mitigation", 1)
    add_table(
        doc,
        ["Risk", "Impact", "Likelihood", "Mitigation"],
        [
            ["Low salon adoption", "High", "Medium", "Free onboarding, dedicated partner success team, in-salon QR setup assistance"],
            ["Payment gateway delays", "High", "Low", "Launch with pay-at-counter; add Razorpay when approved"],
            ["Double booking / slot conflicts", "High", "Medium", "Server-side slot locking with optimistic concurrency"],
            ["Poor internet at salons", "Medium", "Medium", "Offline-capable queue mode; SMS fallback notifications"],
            ["Fake reviews", "Medium", "Medium", "Verified booking badge; admin moderation; ML detection (Phase 2)"],
            ["Scalability during peak (festive season)", "High", "Medium", "Load testing; CDN; auto-scaling infrastructure"],
        ],
        col_widths=[1.8, 0.7, 0.9, 2.9],
    )

    # ── 29. OPEN QUESTIONS ──────────────────────────────────────
    add_heading(doc, "29. Open Questions", 1)
    add_numbered(doc, [
        "What is the final platform fee structure — flat fee vs percentage?",
        "Will salon subscription tiers be introduced at launch or post-MVP?",
        "Which SMS/OTP provider will be used for the Indian market?",
        "Should staff members have individual login accounts in Phase 1?",
        "What is the dispute resolution SLA for booking cancellations and refunds?",
        "Will the platform support multi-language (Gujarati) at Ahmedabad launch?",
        "What are the salon verification/KYC requirements for onboarding?",
        "Should backdated bookings be allowed for salon owners only or also admin?",
    ])

    # ── 30. USER STORIES ────────────────────────────────────────
    add_heading(doc, "30. User Stories (Detailed)", 1)
    stories = [
        ("US-C01 — Search for a salon", "As a customer, I want to search salons by name, service, or area so that I can quickly find relevant options near me.",
         ["Given I am on the homepage, when I type 'bridal makeup' and submit, then I see salons offering bridal services sorted by relevance.",
          "Given I click a popular search chip, then the salon listing page opens with that filter applied."]),
        ("US-C02 — View salon details before booking", "As a customer, I want to see photos, services, staff, reviews, and location so that I can make an informed decision.",
         ["Given I open a salon detail page, when I scroll through sections, then the navigation highlights the active section.",
          "Given I select a service, then the booking panel reflects my selection with price and duration."]),
        ("US-C03 — Book an appointment online", "As a customer, I want to book a specific service at a chosen date and time so that I don't have to call the salon.",
         ["Given I am logged in with service, date, and time selected, when I confirm payment, then I receive a booking confirmation with unique ID.",
          "Given I am not logged in, when I attempt to book, then the login modal appears without losing my selections."]),
        ("US-C04 — Pay online or at counter", "As a customer, I want to choose how I pay so that I have flexibility based on my preference.",
         ["Given I select 'Pay Online', when payment succeeds via Razorpay, then booking status becomes 'confirmed'.",
          "Given I select 'Pay at Salon', when I confirm, then booking is created with 'pay at counter' flag."]),
        ("US-C05 — Track my booking history", "As a customer, I want to view past and upcoming bookings so that I can manage my appointments.",
         ["Given I have upcoming bookings, when I open my profile, then I see them listed with date, salon, and status.",
          "Given a booking is more than 2 hours away, when I cancel, then it is cancelled without penalty."]),
        ("US-C06 — Join walk-in queue", "As a walk-in customer, I want to scan a QR code and join the queue so that I don't wait standing at the reception.",
         ["Given I scan the salon QR, when I select a service and join, then I see my queue position and estimated wait.",
          "Given my turn is next, when the salon advances the queue, then I receive a notification."]),
        ("US-S01 — Manage live queue", "As a salon owner, I want to see and advance my live queue so that I can serve customers in order.",
         ["Given customers are waiting, when I click 'Call Next', then the first waiting customer moves to in-progress.",
          "Given the queue is empty, when I view the queue tab, then I see an empty state with option to add walk-in."]),
        ("US-S02 — View daily revenue", "As a salon owner, I want to see today's earnings and booking count so that I can track business performance.",
         ["Given I open the dashboard overview, when the page loads, then KPI cards show today's bookings, revenue, and queue count.",
          "Given I switch to analytics, then I see earnings trend and peak hour charts."]),
        ("US-A01 — Approve new salon", "As a platform admin, I want to review and approve salon registrations so that only verified salons appear on the marketplace.",
         ["Given a salon registration is pending, when I approve it, then the salon status changes to 'active' and appears in search.",
          "Given I reject a registration, then the owner is notified with reason."]),
    ]
    for title, story, criteria in stories:
        add_heading(doc, title, 2)
        add_para(doc, story, italic=True)
        add_para(doc, "Acceptance Criteria:", bold=True)
        add_bullets(doc, criteria)

    # ── 31. ACCEPTANCE CRITERIA ─────────────────────────────────
    add_heading(doc, "31. Acceptance Criteria — Feature Sign-Off", 1)
    add_para(doc, "The following criteria must be met for each major feature to be considered complete for MVP launch.")
    add_table(
        doc,
        ["Feature", "Acceptance Criteria", "Verified By"],
        [
            ["Homepage search", "Search returns relevant results in <1s; popular chips work; mobile responsive", "QA + Product"],
            ["Salon listing filters", "All filter combinations work; active filters display; clear all resets state", "QA"],
            ["Salon detail booking panel", "Calendar, slots, staff selection work; panel scrolls; toast on validation error", "QA + Product"],
            ["Booking flow", "All 4 steps complete; express flow from salon page lands on payment; summary accurate", "QA + Product"],
            ["Payment", "Razorpay sandbox transactions succeed; failure shows toast; counter payment creates booking", "QA + Engineering"],
            ["Auth", "Login/register persists session; protected routes redirect; ?auth=login opens modal", "QA"],
            ["Profile", "Booking history accurate; cancel works; wishlist add/remove persists", "QA"],
            ["Salon dashboard", "Queue advance works; walk-in entry adds to queue; KPIs match API data", "QA + Partner UAT"],
            ["Admin panel", "Salon approve/reject updates status; user suspend works; analytics load", "QA + Admin UAT"],
            ["Notifications", "Toast appears bottom-right; auto-dismiss 3.5s; success/warning/error styles correct", "QA"],
        ],
        col_widths=[1.4, 3.6, 1.5],
    )

    # ── 32. ERROR HANDLING ──────────────────────────────────────
    add_heading(doc, "32. Error Handling & Edge Cases", 1)
    add_table(
        doc,
        ["Scenario", "Expected Behavior", "User Message"],
        [
            ["No service selected before booking", "Block action; show warning toast", "Please select a service first"],
            ["No date/time selected", "Block action; show warning toast", "Please select a date and time slot"],
            ["Slot becomes unavailable during checkout", "Refresh slots; prompt re-selection", "This slot is no longer available. Please choose another."],
            ["Payment fails (Razorpay)", "Keep booking in draft; show error toast", "Payment failed. Please try again."],
            ["Network error during API call", "Show retry option; preserve form state", "Something went wrong. Please check your connection."],
            ["Session expired during booking", "Redirect to login; restore state after auth", "Your session expired. Please sign in to continue."],
            ["Salon closed when booking attempted", "Disable booking CTA; show closed status", "This salon is currently closed."],
            ["Coupon invalid or expired", "Show inline error on coupon field", "Invalid or expired coupon"],
            ["Double queue join attempt", "Prevent duplicate; show info toast", "You are already in the queue at this salon."],
            ["Unauthorized dashboard access", "Redirect to homepage", "You don't have permission to access this page."],
        ],
        col_widths=[1.8, 2.2, 2.5],
    )

    # ── 33. UI/UX GUIDELINES ────────────────────────────────────
    add_heading(doc, "33. UI/UX Design Guidelines", 1)
    add_heading(doc, "33.1 Brand & Visual Identity", 2)
    add_table(
        doc,
        ["Element", "Specification"],
        [
            ["Primary brand name", "SalonBazar (Salon + Bazar)"],
            ["Tagline", "Your Best Look, Perfectly Booked."],
            ["Primary palette", "Charcoal (#1A1714), Gold (#C9A84C), Warm White (#FAF8F5)"],
            ["Typography — Display", "Serif display font for headings and hero text"],
            ["Typography — Body", "Clean sans-serif for body, labels, and UI elements"],
            ["Tone of voice", "Premium, warm, confident, approachable"],
            ["Iconography", "Feather-style line icons (react-icons/fi)"],
        ],
        col_widths=[1.8, 4.7],
    )
    add_heading(doc, "33.2 Interaction Patterns", 2)
    add_bullets(doc, [
        "Toast notifications (bottom-right) for all user feedback — never use browser alert() dialogs",
        "Modal overlay for authentication; does not navigate away from current page",
        "Sticky booking panel on salon detail (desktop); bottom sheet on mobile",
        "Skeleton loading states for all data-fetching views",
        "Smooth scroll navigation between salon detail sections",
        "Progress stepper for multi-step booking flow",
        "Disabled state styling for CTAs when prerequisites not met",
        "Mobile-first responsive breakpoints: 480px, 768px, 1100px",
    ])
    add_heading(doc, "33.3 Accessibility Requirements", 2)
    add_bullets(doc, [
        "All interactive elements must be keyboard accessible",
        "Form fields must have associated labels",
        "Color contrast ratio minimum 4.5:1 for body text",
        "Alt text on all salon and staff images",
        "ARIA labels on icon-only buttons (wishlist, close, menu)",
        "Focus indicators visible on all focusable elements",
    ])

    # ── 34. QA TEST SCENARIOS ───────────────────────────────────
    add_heading(doc, "34. QA Test Scenarios", 1)
    add_heading(doc, "34.1 Customer Booking — Happy Path", 2)
    add_numbered(doc, [
        "Navigate to homepage → search 'haircut' → open first salon result",
        "Select 'Haircut & Styling' service from services section",
        "In booking panel: select tomorrow's date → pick '10:30 AM' slot → select staff (optional)",
        "Click 'Confirm Booking' → sign in if prompted",
        "Verify landing on Payment step (Step 3) with pre-filled summary",
        "Apply coupon 'FIRST50' → verify discount applied",
        "Agree to cancellation policy → select Pay Online → complete Razorpay payment",
        "Verify confirmation screen with booking ID → check profile booking history",
    ])
    add_heading(doc, "34.2 Salon Dashboard — Queue Management", 2)
    add_numbered(doc, [
        "Log in as shop owner → navigate to dashboard",
        "Verify overview KPIs display correctly",
        "Navigate to Live Queue → verify waiting customers listed",
        "Click advance/next → verify first customer moves to in-progress",
        "Add manual walk-in → verify new entry appears in queue",
        "Complete service → verify customer moves to completed state",
    ])
    add_heading(doc, "34.3 Edge Case Tests", 2)
    add_numbered(doc, [
        "Attempt booking without selecting service → verify toast warning",
        "Attempt payment without agreeing to terms → verify toast warning",
        "Simulate Razorpay payment failure → verify error toast and no duplicate booking",
        "Access /booking/:id without login → verify redirect with auth modal",
        "Access /dashboard without shop_owner role → verify access denied",
        "Cancel booking within 2-hour window → verify penalty message (if applicable)",
        "Cancel booking outside 2-hour window → verify free cancellation",
    ])

    # ── 35. ONBOARDING CHECKLIST ─────────────────────────────────
    add_heading(doc, "35. Salon Partner Onboarding Checklist", 1)
    add_para(doc, "The following checklist must be completed before a salon is approved and goes live on SalonBazar.")
    add_table(
        doc,
        ["Step", "Task", "Owner", "Required"],
        [
            ["1", "Salon owner submits registration form with business details", "Salon Owner", "Yes"],
            ["2", "Upload salon photos (minimum 4 interior/exterior images)", "Salon Owner", "Yes"],
            ["3", "Define service catalog with prices and durations", "Salon Owner", "Yes"],
            ["4", "Add staff profiles with photos and specialties", "Salon Owner", "Yes"],
            ["5", "Set working hours for each day of the week", "Salon Owner", "Yes"],
            ["6", "Provide business address with GPS coordinates", "Salon Owner", "Yes"],
            ["7", "Admin verifies business license / GST (KYC)", "Admin", "Yes"],
            ["8", "Admin reviews photo quality and service pricing", "Admin", "Yes"],
            ["9", "Configure Razorpay sub-merchant account (if online payments)", "Engineering", "Optional"],
            ["10", "Generate and print QR code for walk-in queue", "Salon Owner", "Yes"],
            ["11", "Train salon staff on dashboard usage (30-min session)", "Partner Success", "Yes"],
            ["12", "Admin approves salon → status set to 'active'", "Admin", "Yes"],
        ],
        col_widths=[0.5, 3.5, 1.2, 0.8],
    )

    # ── 36. GLOSSARY ──────────────────────────────────────────────
    add_heading(doc, "36. Glossary & Appendices", 1)
    add_heading(doc, "36.1 Glossary", 2)
    add_table(
        doc,
        ["Term", "Definition"],
        [
            ["GMV", "Gross Merchandise Value — total booking value processed on platform"],
            ["Walk-In", "Customer who arrives at salon without pre-booking; joins live queue"],
            ["Queue Position", "Customer's place in the live waiting line at a salon"],
            ["Platform Fee", "Service charge added to each online booking paid to SalonBazar"],
            ["Express Booking", "Booking initiated from salon detail with pre-selected service/date/time"],
            ["Shop Owner", "Salon business owner with dashboard access"],
            ["Verified Review", "Review from a customer with a confirmed completed booking"],
            ["Slot", "Available time window for a service appointment"],
            ["KYC", "Know Your Customer — identity verification for salon partners"],
        ],
        col_widths=[1.5, 5.0],
    )

    add_heading(doc, "36.2 Appendix A — Valid Coupon Codes (Launch)", 2)
    add_table(
        doc,
        ["Code", "Discount", "Conditions"],
        [
            ["FIRST50", "₹50 off", "First booking for new users"],
            ["SALON20", "20% off (up to ₹200)", "Any booking"],
            ["BAZAR100", "₹100 off", "Minimum order value ₹500"],
        ],
        col_widths=[1.2, 1.8, 3.5],
    )

    add_heading(doc, "36.3 Appendix B — Supported Service Categories", 2)
    add_bullets(doc, [
        "Hair (cut, color, styling, keratin, spa, blow-dry)",
        "Skin & Facial (classic, gold, anti-aging, cleanup)",
        "Bridal (makeup, hair styling, pre-bridal packages)",
        "Nails (manicure, gel extensions, nail art)",
        "Spa & Massage (Swedish, head & shoulder, body polishing)",
        "Men's Grooming (fade cuts, beard shaping, hair color)",
    ])

    add_heading(doc, "36.4 Appendix C — Contact Information", 2)
    add_table(
        doc,
        ["Channel", "Contact"],
        [
            ["Website", "https://salonbazar.shop"],
            ["API", "https://api.salonbazar.shop"],
            ["Support Email", "support@salonbazar.shop"],
            ["Business Email", "hello@salonbazar.shop"],
        ],
        col_widths=[1.5, 5.0],
    )

    add_heading(doc, "36.5 Appendix D — Document End", 2)
    add_para(
        doc,
        "— End of SalonBazar Product Requirements Document —",
        bold=True,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"© {date.today().year} SalonBazar. All rights reserved.")
    run.font.size = Pt(10)
    run.font.italic = True

    doc.save(OUTPUT)
    print(f"PRD saved to {OUTPUT}")


if __name__ == "__main__":
    build_document()
